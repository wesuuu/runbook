"""Unit tests for template_converter service helpers.

Tests the pure/mechanical functions that don't require an AI model:
- _extract_jinja_variables: extracting variable names from text
- _try_render: rendering a template with mock data and checking for issues
- _apply_substitutions_to_docx: find-and-replace on DOCX files
- _to_pdf: converting input files to PDF
- ConversionState: filesystem-based conversion session management
- EventStream: buffered async event stream for SSE
"""

import asyncio
import json
import tempfile
from io import BytesIO
from pathlib import Path
from uuid import uuid4

import pytest
from docx import Document

from app.services.protocols.template_converter import (
    JINJA_PATTERN,
    ConversionState,
    EventStream,
    _apply_substitutions_to_docx,
    _extract_jinja_variables,
    _to_pdf,
    _try_render,
)

# ── _extract_jinja_variables tests ──


class TestExtractJinjaVariables:
    def test_extracts_simple_variables(self):
        text = "Name: {{ protocol_name }}, By: {{ operator_name }}"
        result = _extract_jinja_variables(text)
        assert "protocol_name" in result
        assert "operator_name" in result

    def test_extracts_dotted_variable_top_level(self):
        text = "{{ step.name }} {{ step.description }}"
        result = _extract_jinja_variables(text)
        assert "step" in result

    def test_extracts_loop_collections(self):
        text = "{% for step in steps %}{{ step.name }}{% endfor %}"
        result = _extract_jinja_variables(text)
        assert "steps" in result

    def test_empty_input(self):
        result = _extract_jinja_variables("Plain text")
        assert result == set()

    def test_deduplicates(self):
        text = "{{ name }} and {{ name }}"
        result = _extract_jinja_variables(text)
        assert result == {"name"}


# ── JINJA_PATTERN tests ──


class TestJinjaPattern:
    def test_matches_double_braces(self):
        assert JINJA_PATTERN.findall("Hello {{ name }} world")

    def test_matches_block_tags(self):
        assert JINJA_PATTERN.findall("{% for x in y %}")

    def test_no_match_on_plain_text(self):
        assert not JINJA_PATTERN.findall("Hello world")


# ── _apply_substitutions_to_docx tests ──


class TestApplySubstitutions:
    def _make_docx(self, *paragraphs: str) -> bytes:
        doc = Document()
        for text in paragraphs:
            doc.add_paragraph(text)
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _make_docx_with_table(self) -> bytes:
        doc = Document()
        doc.add_paragraph("Header text")
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Name:"
        table.rows[0].cells[1].text = "Dr. Sarah Chen"
        table.rows[1].cells[0].text = "Date:"
        table.rows[1].cells[1].text = "2026-01-15"
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def test_simple_substitution(self):
        """Should replace text in paragraphs."""
        docx = self._make_docx("Prepared by: Dr. Sarah Chen")
        result, matched, unmatched = _apply_substitutions_to_docx(
            docx,
            [
                {"find": "Dr. Sarah Chen", "replace": "{{ operator_name }}"},
            ],
        )
        doc = Document(BytesIO(result))
        assert "{{ operator_name }}" in doc.paragraphs[0].text
        assert len(matched) == 1
        assert len(unmatched) == 0

    def test_multiple_substitutions(self):
        """Should apply multiple substitutions."""
        docx = self._make_docx(
            "Name: John Smith",
            "Date: 2026-01-15",
        )
        result, matched, _ = _apply_substitutions_to_docx(
            docx,
            [
                {"find": "John Smith", "replace": "{{ operator_name }}"},
                {"find": "2026-01-15", "replace": "{{ effective_date }}"},
            ],
        )
        doc = Document(BytesIO(result))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "{{ operator_name }}" in text
        assert "{{ effective_date }}" in text
        assert len(matched) == 2

    def test_substitution_in_table_cells(self):
        """Should replace text inside table cells."""
        docx = self._make_docx_with_table()
        result, matched, _ = _apply_substitutions_to_docx(
            docx,
            [
                {"find": "Dr. Sarah Chen", "replace": "{{ operator_name }}"},
                {"find": "2026-01-15", "replace": "{{ date }}"},
            ],
        )
        doc = Document(BytesIO(result))
        all_text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text + " "
        assert "{{ operator_name }}" in all_text
        assert "{{ date }}" in all_text
        assert len(matched) == 2

    def test_preserves_unmatched_text(self):
        """Text that doesn't match should be unchanged."""
        docx = self._make_docx("Static label: some value")
        result, _, _ = _apply_substitutions_to_docx(
            docx,
            [
                {"find": "some value", "replace": "{{ var }}"},
            ],
        )
        doc = Document(BytesIO(result))
        assert "Static label:" in doc.paragraphs[0].text

    def test_produces_valid_docx(self):
        """Output should always be a valid DOCX."""
        docx = self._make_docx("Test document")
        result, _, _ = _apply_substitutions_to_docx(
            docx,
            [
                {"find": "Test", "replace": "{{ test_var }}"},
            ],
        )
        assert result[:2] == b"PK"
        doc = Document(BytesIO(result))
        assert len(doc.paragraphs) > 0

    def test_no_match_reports_unmatched(self):
        """Unmatched finds should be reported."""
        docx = self._make_docx("Hello world")
        result, matched, unmatched = _apply_substitutions_to_docx(
            docx,
            [
                {"find": "nonexistent text", "replace": "{{ var }}"},
            ],
        )
        doc = Document(BytesIO(result))
        assert doc.paragraphs[0].text == "Hello world"
        assert len(matched) == 0
        assert "nonexistent text" in unmatched


# ── ConversionState tests ──


class TestConversionState:
    def test_ensure_dir_creates_directory(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            org_id = uuid4()
            conv_id = uuid4()
            state = ConversionState(org_id, conv_id, storage_root=tmpdir)
            state.ensure_dir()
            expected_dir = (
                Path(tmpdir) / str(org_id) / "tmp" / "conversions" / str(conv_id)
            )
            assert expected_dir.is_dir()

    def test_write_and_read(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            org_id = uuid4()
            conv_id = uuid4()
            state = ConversionState(org_id, conv_id, storage_root=tmpdir)
            state.ensure_dir()
            state.write("test.txt", b"hello")
            assert state.read("test.txt") == b"hello"

    def test_exists(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            org_id = uuid4()
            conv_id = uuid4()
            state = ConversionState(org_id, conv_id, storage_root=tmpdir)
            state.ensure_dir()
            assert not state.exists("missing.txt")
            state.write("present.txt", b"data")
            assert state.exists("present.txt")

    def test_write_json_and_read_json(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            org_id = uuid4()
            conv_id = uuid4()
            state = ConversionState(org_id, conv_id, storage_root=tmpdir)
            state.ensure_dir()
            data = {"vars": ["protocol_name"], "count": 1}
            state.write_json("meta.json", data)
            result = state.read_json("meta.json")
            assert result == data

    def test_preview_url(self):
        conv_id = uuid4()
        state = ConversionState(uuid4(), conv_id)
        assert state.preview_url == (
            f"/science/templates/conversions/{conv_id}/preview.pdf"
        )

    def test_template_url(self):
        conv_id = uuid4()
        state = ConversionState(uuid4(), conv_id)
        assert state.template_url == (
            f"/science/templates/conversions/{conv_id}/template.docx"
        )


# ── _to_pdf tests ──


class TestToPdf:
    @pytest.mark.asyncio
    async def test_pdf_passthrough(self):
        fake_pdf = b"%PDF-1.4 fake content"
        result = await _to_pdf(fake_pdf, "document.pdf")
        assert result == fake_pdf

    @pytest.mark.asyncio
    async def test_image_passthrough(self):
        fake_img = b"\x89PNG fake image"
        for ext in ("image.png", "photo.jpg", "scan.jpeg"):
            result = await _to_pdf(fake_img, ext)
            assert result == fake_img


# ── _try_render tests ──


class TestTryRender:
    def test_valid_template_renders(self):
        doc = Document()
        doc.add_paragraph("Protocol: {{ protocol_name }}")
        buf = BytesIO()
        doc.save(buf)
        result = _try_render(buf.getvalue(), "SOP")
        assert result.jinja_remnants == []

    def test_plain_text_template(self):
        doc = Document()
        doc.add_paragraph("This is plain text with no variables.")
        buf = BytesIO()
        doc.save(buf)
        result = _try_render(buf.getvalue(), "SOP")
        assert result.jinja_remnants == []
        assert not result.render_error


# ── EventStream tests ──


class TestEventStream:
    @pytest.mark.asyncio
    async def test_push_and_iterate(self):
        stream = EventStream()
        stream.push("tool_call", {"tool": "apply_substitutions"})
        stream.push("tool_result", {"tool": "apply_substitutions"})
        stream.close()

        events = []
        async for event_type, data_json in stream.iter_events():
            events.append((event_type, json.loads(data_json)))

        assert len(events) == 2
        assert events[0][0] == "tool_call"
        assert events[1][0] == "tool_result"

    @pytest.mark.asyncio
    async def test_late_join_replay(self):
        stream = EventStream()
        stream.push("tool_call", {"tool": "validate"})
        stream.push("complete", {"template_url": "/t.docx"})
        stream.close()

        events = []
        async for event_type, _ in stream.iter_events():
            events.append(event_type)
        assert events == ["tool_call", "complete"]

    @pytest.mark.asyncio
    async def test_close_terminates_iterator(self):
        stream = EventStream()

        async def producer():
            await asyncio.sleep(0.05)
            stream.push("tool_call", {"tool": "validate"})
            await asyncio.sleep(0.05)
            stream.close()

        events = []

        async def consumer():
            async for event_type, _ in stream.iter_events():
                events.append(event_type)

        await asyncio.gather(producer(), consumer())
        assert events == ["tool_call"]
