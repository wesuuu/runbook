"""QA-0008 permutation suite.

Each Pn renders against the templates it's configured for, then asserts
that expected_on substrings appear and expected_off substrings do not.

Pass --write-artifacts to also emit .docx + .pdf into
tests/fixtures/template-permutations/rendered/<Pn>/.
"""

from __future__ import annotations

import io
import subprocess
from pathlib import Path
from unittest.mock import patch

import pytest
from docx import Document

from app.models.protocols import Protocol
from app.models.templates import DocumentTemplate
from app.services.protocols.template_engine import build_context, render_to_docx
from tests.integration.fixtures.template_permutations import builders

SOP_PATH = (
    Path(__file__).resolve().parents[2]
    / "app/services/documents/templates/sop_default.docx"
)
BR_PATH = (
    Path(__file__).resolve().parents[2]
    / "app/services/documents/templates/batch_record_default.docx"
)
ARTIFACT_ROOT = (
    Path(__file__).resolve().parents[3]
    / "tests/fixtures/template-permutations/rendered"
)


def _doc_text(blob: bytes) -> str:
    d = Document(io.BytesIO(blob))
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _convert_to_pdf(docx_path: Path) -> Path | None:
    try:
        subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(docx_path.parent),
                str(docx_path),
            ],
            check=True,
            capture_output=True,
            timeout=60,
        )
    except (
        FileNotFoundError,
        subprocess.CalledProcessError,
        subprocess.TimeoutExpired,
    ):
        return None
    pdf_path = docx_path.with_suffix(".pdf")
    return pdf_path if pdf_path.exists() else None


@pytest.mark.parametrize(
    "builder_name,template_key,template_path",
    [
        ("build_p1", "sop", str(SOP_PATH)),
        ("build_p1", "batch_record", str(BR_PATH)),
        ("build_p2", "sop", str(SOP_PATH)),
        ("build_p3", "sop", str(SOP_PATH)),
        ("build_p4", "batch_record", str(BR_PATH)),
        ("build_p5", "batch_record", str(BR_PATH)),
        ("build_p6", "batch_record", str(BR_PATH)),
    ],
)
def test_permutation_renders(
    builder_name, template_key, template_path, write_artifacts
):
    built = getattr(builders, builder_name)()
    if template_key not in built.renders_against:
        pytest.skip(f"{built.name} not configured to render against {template_key}")

    ctx, unresolved = build_context(**built.kwargs)
    ctx.setdefault("approval", None)
    ctx.setdefault("approval_history", [])
    ctx.setdefault("unapproved_warning", "")
    if built.context_overrides:
        ctx.update(built.context_overrides)

    assert unresolved == [], f"{built.name}: unresolved tokens: {unresolved}"

    docx_bytes = render_to_docx(template_path, ctx)
    text = _doc_text(docx_bytes)

    expected_on = built.per_template_expected_on.get(template_key, built.expected_on)
    expected_off = built.per_template_expected_off.get(template_key, built.expected_off)
    for needle in expected_on:
        assert needle in text, f"{built.name}/{template_key}: missing '{needle}'"
    for needle in expected_off:
        assert needle not in text, f"{built.name}/{template_key}: unexpected '{needle}'"

    if write_artifacts:
        outdir = ARTIFACT_ROOT / built.name
        outdir.mkdir(parents=True, exist_ok=True)
        docx_path = outdir / f"{template_key}.docx"
        docx_path.write_bytes(docx_bytes)
        _convert_to_pdf(docx_path)


def _graph_from_p1() -> dict:
    """Translate P1's roles_with_steps into the swim-lane graph shape
    that ``_parse_graph_roles_and_steps`` recognises (role lanes wrap
    their member steps via ``parentId``)."""
    built = builders.build_p1()
    kw = built.kwargs
    nodes: list[dict] = []
    edges: list[dict] = []
    for role in kw["roles_with_steps"]:
        lane_id = f"lane-{role['role_name']}"
        nodes.append(
            {
                "id": lane_id,
                "type": "swimLane",
                "data": {"role_name": role["role_name"]},
                "position": {"x": 0, "y": 0},
            }
        )
        for step in role["steps"]:
            nodes.append(
                {
                    "id": step["id"],
                    "type": "unitOp",
                    "parentId": lane_id,
                    "data": {
                        "label": step["name"],
                        "duration_min": step["duration_min"],
                        "params": step["params"],
                        "paramSchema": step["param_schema"],
                        "equipment": step.get("equipment", []),
                    },
                    "position": {"x": 100, "y": 100},
                }
            )
    return {
        "nodes": nodes,
        "edges": edges,
        "timeEnabled": kw.get("time_enabled", False),
        "startTime": kw.get("start_time", ""),
    }


async def test_p1_endpoint_renders_batch_record_pdf(
    client, db_session, test_user, test_org, test_project, auth_headers
):
    """End-to-end: a real Protocol row + the system BR template renders
    a PDF through ``/protocols/{id}/pdf/batch-record`` with no
    leaked Jinja tokens.

    Persists a Protocol carrying P1's kitchen-sink metadata and graph;
    skips Run persistence because the GET endpoint runs preview mode
    (``run_name="Preview"``) and never queries a Run row.
    """
    template = DocumentTemplate(
        name="BR System Default",
        template_type="BATCH_RECORD",
        file_path="system/document_templates/batch_record_default.docx",
        original_filename="batch_record_default.docx",
        mime_type=(
            "application/vnd.openxmlformats-officedocument" ".wordprocessingml.document"
        ),
        file_size_bytes=BR_PATH.stat().st_size,
        org_id=test_org.id,
        is_system=True,
        is_default=True,
    )
    db_session.add(template)
    await db_session.flush()

    built = builders.build_p1()
    kw = built.kwargs
    protocol = Protocol(
        name=kw["protocol_name"],
        description=kw.get("protocol_description", ""),
        project_id=test_project.id,
        version_number=kw.get("version_number", 1),
        graph=_graph_from_p1(),
        batch_record_template_id=template.id,
        doc_number=kw.get("doc_number"),
        purpose=kw.get("purpose"),
        scope=kw.get("scope"),
        references=kw.get("references"),
        definitions=kw.get("definitions"),
        slug="template-permutation-proto",
        owner_org_id=test_project.organization_id,
    )
    db_session.add(protocol)
    await db_session.flush()

    # Storage paths resolve through FileStorageService against the live
    # uploads dir; in the test we substitute the bundled template path
    # so the endpoint reads the system default directly.
    with patch(
        "app.api.endpoints.protocol_pdfs._resolve_template_path",
        return_value=str(BR_PATH),
    ):
        resp = await client.get(
            f"/protocols/{protocol.id}/pdf/batch-record",
            headers=auth_headers,
        )

    assert resp.status_code == 200, resp.text
    assert resp.headers["content-type"].startswith("application/pdf")
    # No unresolved {{token}} should leak through with P1's metadata.
    leak = resp.headers.get("X-Unresolved-Placeholders", "")
    assert leak == "", f"unresolved placeholders leaked: {leak}"
    # PDF body is non-trivial (LibreOffice min-output is several KB).
    assert len(resp.content) > 1000
