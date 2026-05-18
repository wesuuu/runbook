"""Smoke test: new SOP template parses and renders correctly.

Asserts the rendered .docx text contains expected section headings when fed
get_mock_context(), and that no Jinja tokens leak through. Also verifies
truly optional sections (approval, approval history, critical-requirement
callout, unapproved warning) are gated out when the context omits them.
"""

import io
from pathlib import Path

from docx import Document
from docxtpl import DocxTemplate

from app.services.protocols.template_engine import build_context, get_mock_context

SOP_PATH = (
    Path(__file__).resolve().parents[2]
    / "app/services/documents/templates/sop_default.docx"
)


def _doc_text(blob: bytes) -> str:
    parts = []
    d = Document(io.BytesIO(blob))
    for p in d.paragraphs:
        parts.append(p.text)
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_sop_template_renders_mock_context_without_unresolved_tokens():
    tpl = DocxTemplate(SOP_PATH)
    tpl.render(get_mock_context())
    buf = io.BytesIO()
    tpl.save(buf)
    text = _doc_text(buf.getvalue())

    # Section headings introduced by the GLP/bioreactor-style SOP layout.
    for heading in (
        "Standard Operating Procedure",
        "Document Number:",
        "Title:",
        "Effective Date:",
        "1.0 Purpose",
        "2.0 Scope",
        "3.0 Procedure",
        # Role-table column headers
        "Step",
        "Name",
        "Instruction",
        "Duration",
    ):
        assert heading in text, f"missing section: {heading}"

    # No leaked Jinja tokens
    assert "{{" not in text
    assert "{%" not in text


def test_sop_template_omits_optional_sections_when_blank():
    """Sections gated by Jinja conditionals (approval, approval history,
    critical-requirement callout, unapproved warning) must NOT appear when
    the context omits them. Purpose/Scope/Procedure are always rendered
    because they are part of the document skeleton."""
    ctx, _ = build_context(protocol_name="bare")
    ctx.setdefault("approval", None)
    ctx.setdefault("approval_history", [])
    ctx.setdefault("unapproved_warning", "")
    tpl = DocxTemplate(SOP_PATH)
    tpl.render(ctx)
    buf = io.BytesIO()
    tpl.save(buf)
    text = _doc_text(buf.getvalue())

    # Truly optional sections — these should be absent with a bare context.
    assert "Approval & Signatures" not in text
    assert "Approval History" not in text
    assert "CRITICAL REQUIREMENT" not in text
    assert "UNAPPROVED — DRAFT ONLY" not in text

    # Always-present skeleton headings remain.
    assert "1.0 Purpose" in text
    assert "2.0 Scope" in text
    assert "3.0 Procedure" in text
