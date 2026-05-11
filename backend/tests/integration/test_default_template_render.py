"""Smoke test: default SOP template renders the approval block."""

from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.endpoints.protocol_pdfs import _build_approval_context
from app.core.security import hash_password
from app.models.iam import OrganizationMember, User
from app.models.science import Project, Protocol, ProtocolApprovalEvent
from app.services.protocols.template_engine import (build_context,
                                                    render_to_docx)


def _docx_text(docx_bytes: bytes) -> str:
    doc = Document(BytesIO(docx_bytes))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


@pytest.mark.asyncio
async def test_default_sop_renders_with_approval_block(
    db_session: AsyncSession,
    test_org,
    test_project: Project,
    test_user: User,
):
    """Render the default SOP for an APPROVED protocol; output must
    contain the approver's name and an Approval section."""
    approver = User(
        email="default-approver@test.com",
        hashed_password=hash_password("test"),
        full_name="Approver Alice",
        email_verified=True,
    )
    db_session.add(approver)
    await db_session.flush()
    db_session.add(
        OrganizationMember(
            user_id=approver.id,
            organization_id=test_org.id,
            roles=["MEMBER", "PROTOCOL_APPROVER"],
        )
    )

    proto = Protocol(
        name="Default Tpl Proto",
        project_id=test_project.id,
        status="APPROVED",
        version_number=1,
        requires_approval=True,
        approved_by_id=approver.id,
        approved_at=datetime.now(timezone.utc),
        created_by_id=test_user.id,
    )
    db_session.add(proto)
    await db_session.flush()

    db_session.add(
        ProtocolApprovalEvent(
            protocol_id=proto.id,
            actor_id=approver.id,
            action="APPROVED",
            signature_statement="I have reviewed and approved this protocol",
        )
    )
    await db_session.flush()

    approval_ctx = await _build_approval_context(db_session, proto, test_project)

    template_path = (
        Path(__file__).resolve().parent.parent.parent
        / "app/services/documents/templates/sop_default.docx"
    )

    context = build_context(
        protocol_name=proto.name,
        protocol_description="",
        version_number=proto.version_number,
        created_at="January 1, 2026",
        roles_with_steps=[],
        flat_steps=[],
        is_role_based=True,
    )
    context.update(approval_ctx)

    docx_bytes = render_to_docx(template_path, context)
    text = _docx_text(docx_bytes)
    assert "Approval" in text
    assert "Approver Alice" in text
