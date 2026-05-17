"""Smoke test: new Batch Record template parses and renders correctly.

Asserts the rendered .docx text contains expected section headings when fed
get_mock_context(), and that no Jinja tokens leak through. Verifies the
GLP-style numbered sections (1. General Information through 6. Final
Disposition) and the always-present Verifier column appear.
"""
import io
from pathlib import Path

from docx import Document
from docxtpl import DocxTemplate

from app.services.protocols.template_engine import get_mock_context

BR_PATH = (
    Path(__file__).resolve().parents[2]
    / "app/services/documents/templates/batch_record_default.docx"
)


def _doc_text(blob: bytes) -> str:
    parts = []
    d = Document(io.BytesIO(blob))
    for p in d.paragraphs:
        parts.append(p.text)
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_br_template_renders_mock_context_without_unresolved_tokens():
    tpl = DocxTemplate(BR_PATH)
    tpl.render(get_mock_context())
    buf = io.BytesIO()
    tpl.save(buf)
    text = _doc_text(buf.getvalue())

    # GLP-style numbered section headers
    for heading in (
        "Batch Manufacturing Record",
        "1. General Information",
        "2. Bill of Materials (BOM)",
        "3. Equipment Log",
        "4. Execution: Unit Operations",
        "5. Deviations and Process Comments",
        "6. Final Disposition & Signatures",
        "Wet-Ink Sign-Off",
    ):
        assert heading in text, f"missing section: {heading}"

    # No leaked Jinja tokens
    assert "{{" not in text
    assert "{%" not in text

    # Step-execution column headers are always present in BR layout.
    assert "Verifier" in text
    assert "Operator" in text
    assert "Batch / Lot Number" in text
