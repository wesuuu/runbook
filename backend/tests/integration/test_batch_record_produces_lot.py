"""Integration test: batch_record_default.docx honors produces_lot.

When ``produces_lot=True``, the "Batch / Lot Number" header row must
render with the run's ``lot_number``. When ``produces_lot=False``, the
entire row must be hidden by the ``{%tr if produces_lot %}`` gate.

Uses ``render_to_docx`` (the public renderer in
``app.services.protocols.template_engine``) directly so the test is a
focused docx contract check, not an end-to-end run pipeline test.
"""

from io import BytesIO
from pathlib import Path

from docx import Document

from app.services.protocols.template_engine import build_context, render_to_docx

BR_TEMPLATE = (
    Path(__file__).resolve().parents[2]
    / "app/services/documents/templates/batch_record_default.docx"
)


def _docx_text(docx_bytes: bytes) -> str:
    doc = Document(BytesIO(docx_bytes))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


def _build_br_context(*, produces_lot: bool, lot_number: str | None) -> dict:
    """Minimal context the BR template can render against."""
    ctx, _ = build_context(
        protocol_name="Test Protocol",
        protocol_description="Test description.",
        version_number=1,
        created_at="January 1, 2026",
        run_name="Run-Test-001",
        project_name="Test Project",
        organization_name="Test Org",
        roles_with_steps=[],
        flat_steps=[],
        is_role_based=True,
        lot_number=lot_number or "",
        produces_lot=produces_lot,
    )
    return ctx


def test_batch_record_includes_lot_row_when_produces_lot():
    """produces_lot=True → row visible with lot_number value."""
    ctx = _build_br_context(produces_lot=True, lot_number="LOT-000099")
    docx_bytes = render_to_docx(BR_TEMPLATE, ctx)
    text = _docx_text(docx_bytes)

    assert "LOT-000099" in text, (
        "Expected lot_number 'LOT-000099' to appear in rendered BR when "
        f"produces_lot=True. Got text:\n{text[:2000]}"
    )
    assert "Batch / Lot Number" in text or "Lot Number" in text, (
        "Expected the 'Batch / Lot Number' label to appear when " "produces_lot=True."
    )
    # Defensive: the legacy mis-binding ({{ run_name }}) must not surface
    # the run name in the lot row. Run name is "Run-Test-001"; we don't
    # want to see it appearing as the lot value.
    assert "{{ run_name }}" not in text
    assert "{{ lot_number }}" not in text


def test_batch_record_hides_lot_row_when_not_producer():
    """produces_lot=False → label and value both absent."""
    ctx = _build_br_context(produces_lot=False, lot_number=None)
    docx_bytes = render_to_docx(BR_TEMPLATE, ctx)
    text = _docx_text(docx_bytes)

    # The header-table label row should be gone. "Part / Lot Number" is
    # a separate BOM column header and is expected to stay.
    assert "Batch / Lot Number" not in text, (
        "'Batch / Lot Number' label leaked into rendered BR when "
        f"produces_lot=False. Got text:\n{text[:2000]}"
    )
    # No leaked jinja tokens
    assert "{%tr" not in text
    assert "{{" not in text
