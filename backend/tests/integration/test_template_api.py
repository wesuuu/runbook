"""Integration tests for the template management API (Phase 1)."""

from io import BytesIO
from uuid import UUID

import pytest
import pytest_asyncio
from docx import Document

from app.core.security import create_access_token, hash_password
from app.models.iam import Organization, OrganizationMember, User
from app.models.science import Project, Protocol
from app.models.templates import DocumentTemplate

# ── Helpers ──


def _make_test_docx() -> bytes:
    """Create a minimal .docx with Jinja2 template variables."""
    doc = Document()
    doc.add_paragraph("{{ protocol_name }}")
    doc.add_paragraph("{% for step in steps %}{{ step.name }}{% endfor %}")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


DOCX_MIME = "application/vnd.openxmlformats-officedocument" ".wordprocessingml.document"


# ── Fixtures ──


@pytest_asyncio.fixture
async def member_user(db_session, test_org) -> User:
    """Non-admin member of test_org."""
    user = User(
        email="member@example.com",
        hashed_password=hash_password("testpass"),
        full_name="Regular Member",
        selected_org_id=test_org.id,
        email_verified=True,
    )
    db_session.add(user)
    await db_session.flush()
    db_session.add(
        OrganizationMember(
            user_id=user.id,
            organization_id=test_org.id,
            roles=["MEMBER"],
        )
    )
    await db_session.flush()
    return user


@pytest_asyncio.fixture
async def member_headers(member_user, test_org) -> dict:
    token = create_access_token(
        member_user.id,
        org_id=test_org.id,
        subscription_tier=test_org.subscription_tier,
        email_verified=True,
    )
    return {"Authorization": f"Bearer {token}"}


@pytest_asyncio.fixture
async def test_project(db_session, test_org) -> Project:
    project = Project(
        name="Test Project",
        organization_id=test_org.id,
    )
    db_session.add(project)
    await db_session.flush()
    return project


@pytest_asyncio.fixture
async def system_templates(db_session) -> dict[str, UUID]:
    """Ensure system templates exist and return their IDs."""
    from sqlalchemy import select

    sop_result = await db_session.execute(
        select(DocumentTemplate.id).where(
            DocumentTemplate.is_system == True,
            DocumentTemplate.is_default == True,
            DocumentTemplate.template_type == "SOP",
        )
    )
    sop_id = sop_result.scalar_one_or_none()

    br_result = await db_session.execute(
        select(DocumentTemplate.id).where(
            DocumentTemplate.is_system == True,
            DocumentTemplate.is_default == True,
            DocumentTemplate.template_type == "BATCH_RECORD",
        )
    )
    br_id = br_result.scalar_one_or_none()

    # If not seeded (test DB), create them
    if not sop_id:
        sop = DocumentTemplate(
            name="System SOP",
            template_type="SOP",
            file_path="system/document_templates/sop_default.docx",
            original_filename="sop_default.docx",
            mime_type=DOCX_MIME,
            file_size_bytes=0,
            is_system=True,
            is_default=True,
        )
        db_session.add(sop)
        await db_session.flush()
        sop_id = sop.id

    if not br_id:
        br = DocumentTemplate(
            name="System BR",
            template_type="BATCH_RECORD",
            file_path="system/document_templates/batch_record_default.docx",
            original_filename="batch_record_default.docx",
            mime_type=DOCX_MIME,
            file_size_bytes=0,
            is_system=True,
            is_default=True,
        )
        db_session.add(br)
        await db_session.flush()
        br_id = br.id

    return {"SOP": sop_id, "BATCH_RECORD": br_id}


# ── Preview tests ──


async def test_preview_returns_pdf_and_variables(client, auth_headers):
    docx_bytes = _make_test_docx()
    resp = await client.post(
        "/templates/preview",
        files={"file": ("test.docx", docx_bytes, DOCX_MIME)},
        data={"template_type": "SOP"},
        headers=auth_headers,
    )
    assert resp.status_code == 200
    assert resp.headers["content-type"] == "application/pdf"
    assert resp.content[:4] == b"%PDF"
    assert "X-Recognized-Variables" in resp.headers
    assert "X-Unrecognized-Variables" in resp.headers


async def test_preview_rejects_non_docx(client, auth_headers):
    resp = await client.post(
        "/templates/preview",
        files={"file": ("test.txt", b"not a docx", "text/plain")},
        data={"template_type": "SOP"},
        headers=auth_headers,
    )
    assert resp.status_code == 422


async def test_preview_rejects_non_admin(client, member_headers):
    docx_bytes = _make_test_docx()
    resp = await client.post(
        "/templates/preview",
        files={"file": ("test.docx", docx_bytes, DOCX_MIME)},
        data={"template_type": "SOP"},
        headers=member_headers,
    )
    assert resp.status_code == 403


# ── Create tests ──


async def test_create_template_stores_and_returns(client, auth_headers):
    docx_bytes = _make_test_docx()
    resp = await client.post(
        "/templates",
        files={"file": ("my_sop.docx", docx_bytes, DOCX_MIME)},
        data={"name": "My SOP Template", "template_type": "SOP"},
        headers=auth_headers,
    )
    assert resp.status_code == 200
    data = resp.json()
    assert data["name"] == "My SOP Template"
    assert data["template_type"] == "SOP"
    assert data["original_filename"] == "my_sop.docx"
    assert data["status"] == "ACTIVE"
    assert data["is_system"] is False


async def test_create_duplicate_filename_rejected(client, auth_headers):
    docx_bytes = _make_test_docx()
    # First upload
    resp1 = await client.post(
        "/templates",
        files={"file": ("duplicate.docx", docx_bytes, DOCX_MIME)},
        data={"name": "First", "template_type": "SOP"},
        headers=auth_headers,
    )
    assert resp1.status_code == 200

    # Second upload with same filename
    resp2 = await client.post(
        "/templates",
        files={"file": ("duplicate.docx", docx_bytes, DOCX_MIME)},
        data={"name": "Second", "template_type": "SOP"},
        headers=auth_headers,
    )
    assert resp2.status_code == 409


async def test_create_with_set_as_default(client, auth_headers, test_org, db_session):
    docx_bytes = _make_test_docx()
    resp = await client.post(
        "/templates",
        files={"file": ("default_sop.docx", docx_bytes, DOCX_MIME)},
        data={
            "name": "Default SOP",
            "template_type": "SOP",
            "set_as_default": "true",
        },
        headers=auth_headers,
    )
    assert resp.status_code == 200
    data = resp.json()
    assert data["is_current_default"] is True

    # Verify org was updated
    from sqlalchemy import select

    result = await db_session.execute(
        select(Organization).where(Organization.id == test_org.id)
    )
    org = result.scalar_one()
    assert str(org.default_sop_template_id) == data["id"]


async def test_non_admin_cannot_create(client, member_headers):
    docx_bytes = _make_test_docx()
    resp = await client.post(
        "/templates",
        files={"file": ("test.docx", docx_bytes, DOCX_MIME)},
        data={"name": "Should Fail", "template_type": "SOP"},
        headers=member_headers,
    )
    assert resp.status_code == 403


# ── List tests ──


async def test_list_templates(client, auth_headers):
    # Create two templates
    docx_bytes = _make_test_docx()
    for name, fname in [("SOP A", "sop_a.docx"), ("BR B", "br_b.docx")]:
        ttype = "SOP" if "SOP" in name else "BATCH_RECORD"
        await client.post(
            "/templates",
            files={"file": (fname, docx_bytes, DOCX_MIME)},
            data={"name": name, "template_type": ttype},
            headers=auth_headers,
        )

    resp = await client.get("/templates", headers=auth_headers)
    assert resp.status_code == 200
    data = resp.json()
    # Should include user templates + system templates
    names = [t["name"] for t in data]
    assert "SOP A" in names
    assert "BR B" in names


async def test_list_filters_by_type(client, auth_headers):
    docx_bytes = _make_test_docx()
    await client.post(
        "/templates",
        files={"file": ("filter_test.docx", docx_bytes, DOCX_MIME)},
        data={"name": "Filter SOP", "template_type": "SOP"},
        headers=auth_headers,
    )

    resp = await client.get(
        "/templates?template_type=BATCH_RECORD", headers=auth_headers
    )
    assert resp.status_code == 200
    names = [t["name"] for t in resp.json()]
    assert "Filter SOP" not in names


async def test_list_includes_system_templates(client, auth_headers, system_templates):
    resp = await client.get("/templates", headers=auth_headers)
    assert resp.status_code == 200
    system_names = [t["name"] for t in resp.json() if t["is_system"]]
    assert len(system_names) >= 1


# ── Detail tests ──


async def test_get_template_detail(client, auth_headers):
    docx_bytes = _make_test_docx()
    create_resp = await client.post(
        "/templates",
        files={"file": ("detail_test.docx", docx_bytes, DOCX_MIME)},
        data={"name": "Detail Test", "template_type": "SOP"},
        headers=auth_headers,
    )
    template_id = create_resp.json()["id"]

    resp = await client.get(f"/templates/{template_id}", headers=auth_headers)
    assert resp.status_code == 200
    assert resp.json()["name"] == "Detail Test"


# ── Update tests ──


async def test_update_name_and_description(client, auth_headers):
    docx_bytes = _make_test_docx()
    create_resp = await client.post(
        "/templates",
        files={"file": ("update_test.docx", docx_bytes, DOCX_MIME)},
        data={"name": "Original Name", "template_type": "SOP"},
        headers=auth_headers,
    )
    template_id = create_resp.json()["id"]

    resp = await client.put(
        f"/templates/{template_id}",
        json={"name": "Updated Name", "description": "New desc"},
        headers=auth_headers,
    )
    assert resp.status_code == 200
    assert resp.json()["name"] == "Updated Name"
    assert resp.json()["description"] == "New desc"


async def test_archive_template(client, auth_headers):
    docx_bytes = _make_test_docx()
    create_resp = await client.post(
        "/templates",
        files={"file": ("archive_test.docx", docx_bytes, DOCX_MIME)},
        data={"name": "To Archive", "template_type": "SOP"},
        headers=auth_headers,
    )
    template_id = create_resp.json()["id"]

    resp = await client.put(
        f"/templates/{template_id}",
        json={"status": "ARCHIVED"},
        headers=auth_headers,
    )
    assert resp.status_code == 200
    assert resp.json()["status"] == "ARCHIVED"
    assert resp.json()["archived_at"] is not None


async def test_unarchive_template(client, auth_headers):
    docx_bytes = _make_test_docx()
    create_resp = await client.post(
        "/templates",
        files={"file": ("unarchive_test.docx", docx_bytes, DOCX_MIME)},
        data={"name": "To Unarchive", "template_type": "SOP"},
        headers=auth_headers,
    )
    template_id = create_resp.json()["id"]

    # Archive
    await client.put(
        f"/templates/{template_id}",
        json={"status": "ARCHIVED"},
        headers=auth_headers,
    )

    # Unarchive
    resp = await client.put(
        f"/templates/{template_id}",
        json={"status": "ACTIVE"},
        headers=auth_headers,
    )
    assert resp.status_code == 200
    assert resp.json()["status"] == "ACTIVE"
    assert resp.json()["archived_at"] is None


async def test_system_template_immutable(client, auth_headers, system_templates):
    sop_id = str(system_templates["SOP"])
    resp = await client.put(
        f"/templates/{sop_id}",
        json={"name": "Hacked"},
        headers=auth_headers,
    )
    assert resp.status_code == 403


# ── Variables reference ──


async def test_get_variables_reference(client, auth_headers):
    resp = await client.get("/templates/variables", headers=auth_headers)
    assert resp.status_code == 200
    data = resp.json()
    assert "protocol" in data
    assert "run" in data
    assert "loops" in data
    assert "step_fields" in data
    assert "special" in data


# ── Default resolution on protocol creation ──


async def test_protocol_creation_stamps_defaults(
    client, auth_headers, test_project, system_templates, db_session
):
    """New protocol gets system default template IDs."""
    # Set system defaults on the org first
    from sqlalchemy import select

    result = await db_session.execute(
        select(Organization).where(Organization.id == test_project.organization_id)
    )
    org = result.scalar_one()
    org.default_sop_template_id = system_templates["SOP"]
    org.default_batch_record_template_id = system_templates["BATCH_RECORD"]
    await db_session.flush()

    resp = await client.post(
        "/science/protocols",
        json={
            "name": "Test Protocol",
            "project_id": str(test_project.id),
            "graph": {"nodes": [], "edges": []},
        },
        headers=auth_headers,
    )
    assert resp.status_code in (200, 201)

    # Verify template IDs were stamped
    proto_id = resp.json()["id"]
    result = await db_session.execute(select(Protocol).where(Protocol.id == proto_id))
    proto = result.scalar_one()
    assert proto.sop_template_id == system_templates["SOP"]
    assert proto.batch_record_template_id == system_templates["BATCH_RECORD"]
