"""Integration tests for template conversion API endpoints.

Tests the HTTP layer: file upload, validation, 404 handling, file serving.
AI-dependent tests (convert, refine) are mocked to avoid real LLM calls.
"""

from io import BytesIO
from unittest.mock import AsyncMock, patch
from uuid import uuid4

import pytest
from docx import Document

DOCX_MIME = "application/vnd.openxmlformats-officedocument" ".wordprocessingml.document"


# ── Helpers ──


def _make_filled_docx() -> bytes:
    """Create a minimal filled .docx (simulating a completed SOP)."""
    doc = Document()
    doc.add_heading("Buffer Preparation SOP", level=1)
    doc.add_paragraph("Prepared by: John Smith on 2026-01-15")
    doc.add_paragraph("Lot Number: BR-2026-0042")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_template_docx() -> bytes:
    """Create a minimal Jinja2 template .docx."""
    doc = Document()
    doc.add_heading("{{ protocol_name }}", level=1)
    doc.add_paragraph("Prepared by: {{ operator_name }} on {{ completion_date }}")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Endpoint Tests ──


class TestConvertEndpoint:
    """POST /science/templates/convert"""

    @pytest.mark.asyncio
    async def test_rejects_unauthenticated(self, client):
        """Should return 401 without auth header."""
        resp = await client.post(
            "/science/templates/convert",
            data={"template_type": "SOP"},
            files={"file": ("test.docx", _make_filled_docx(), DOCX_MIME)},
        )
        assert resp.status_code == 401

    @pytest.mark.asyncio
    async def test_rejects_invalid_template_type(self, client, auth_headers):
        """Should return 422 for invalid template_type."""
        resp = await client.post(
            "/science/templates/convert",
            headers=auth_headers,
            data={"template_type": "INVALID"},
            files={"file": ("test.docx", _make_filled_docx(), DOCX_MIME)},
        )
        assert resp.status_code == 422

    @pytest.mark.asyncio
    async def test_rejects_unsupported_file_type(self, client, auth_headers):
        """Should return 422 for unsupported MIME type."""
        resp = await client.post(
            "/science/templates/convert",
            headers=auth_headers,
            data={"template_type": "SOP"},
            files={"file": ("test.txt", b"hello", "text/plain")},
        )
        assert resp.status_code == 422

    @pytest.mark.asyncio
    async def test_successful_convert_starts_async(self, client, auth_headers):
        """Should return 202 with conversion_id for async processing."""
        with patch(
            "app.api.endpoints.template_convert._preflight_ai_check",
            new_callable=AsyncMock,
        ):
            resp = await client.post(
                "/science/templates/convert",
                headers=auth_headers,
                data={"template_type": "SOP"},
                files={
                    "file": (
                        "test.docx",
                        _make_filled_docx(),
                        DOCX_MIME,
                    )
                },
            )
        assert resp.status_code == 202
        data = resp.json()
        assert "conversion_id" in data
        assert data["status"] == "processing"


class TestEventsEndpoint:
    """GET /science/templates/conversions/{id}/events"""

    @pytest.mark.asyncio
    async def test_returns_error_for_missing_conversion(self, client, auth_headers):
        """Should return an SSE error event for a nonexistent conversion."""
        fake_id = uuid4()
        resp = await client.get(
            f"/science/templates/conversions/{fake_id}/events",
            headers=auth_headers,
        )
        assert resp.status_code == 200
        assert "text/event-stream" in resp.headers["content-type"]
        assert "Conversion not found" in resp.text


class TestRefineEndpoint:
    """POST /science/templates/conversions/{id}/refine"""

    @pytest.mark.asyncio
    async def test_returns_404_for_missing_conversion(self, client, auth_headers):
        """Should return 404 if conversion_id doesn't exist."""
        fake_id = uuid4()
        resp = await client.post(
            f"/science/templates/conversions/{fake_id}/refine",
            headers=auth_headers,
            json={"instruction": "rename lot_number to batch_id"},
        )
        assert resp.status_code == 404


class TestReuploadEndpoint:
    """POST /science/templates/conversions/{id}/reupload"""

    @pytest.mark.asyncio
    async def test_returns_404_for_missing_conversion(self, client, auth_headers):
        """Should return 404 if conversion_id doesn't exist."""
        fake_id = uuid4()
        resp = await client.post(
            f"/science/templates/conversions/{fake_id}/reupload",
            headers=auth_headers,
            files={
                "file": (
                    "template.docx",
                    _make_template_docx(),
                    DOCX_MIME,
                )
            },
        )
        assert resp.status_code == 404


class TestSaveEndpoint:
    """POST /science/templates/conversions/{id}/save"""

    @pytest.mark.asyncio
    async def test_returns_404_for_missing_conversion(self, client, auth_headers):
        """Should return 404 if conversion_id doesn't exist."""
        fake_id = uuid4()
        resp = await client.post(
            f"/science/templates/conversions/{fake_id}/save",
            headers=auth_headers,
            json={
                "name": "Test Template",
                "template_type": "SOP",
            },
        )
        assert resp.status_code == 404

    @pytest.mark.asyncio
    async def test_rejects_invalid_template_type(self, client, auth_headers, test_org):
        """Should return 422 for invalid template_type on save."""
        from app.services.protocols.template_converter import ConversionState

        conv_id = uuid4()
        state = ConversionState(test_org.id, conv_id)
        state.ensure_dir()
        state.write("template.docx", _make_template_docx())

        resp = await client.post(
            f"/science/templates/conversions/{conv_id}/save",
            headers=auth_headers,
            json={
                "name": "Test Template",
                "template_type": "INVALID",
            },
        )
        assert resp.status_code == 422


class TestFileServingEndpoints:
    """GET endpoints for serving preview.pdf and template.docx"""

    @pytest.mark.asyncio
    async def test_preview_404_when_missing(self, client, auth_headers):
        """Should return 404 when preview doesn't exist."""
        fake_id = uuid4()
        resp = await client.get(
            f"/science/templates/conversions/{fake_id}/preview.pdf",
            headers=auth_headers,
        )
        assert resp.status_code == 404

    @pytest.mark.asyncio
    async def test_template_404_when_missing(self, client, auth_headers):
        """Should return 404 when template doesn't exist."""
        fake_id = uuid4()
        resp = await client.get(
            f"/science/templates/conversions/{fake_id}/template.docx",
            headers=auth_headers,
        )
        assert resp.status_code == 404

    @pytest.mark.asyncio
    async def test_serves_template_docx(self, client, auth_headers, test_org):
        """Should serve the template file when it exists."""
        from app.services.protocols.template_converter import ConversionState

        conv_id = uuid4()
        state = ConversionState(test_org.id, conv_id)
        state.ensure_dir()
        template_bytes = _make_template_docx()
        state.write("template.docx", template_bytes)

        resp = await client.get(
            f"/science/templates/conversions/{conv_id}/template.docx",
            headers=auth_headers,
        )
        assert resp.status_code == 200
        assert len(resp.content) == len(template_bytes)

    @pytest.mark.asyncio
    async def test_serves_preview_pdf(self, client, auth_headers, test_org):
        """Should serve the preview PDF when it exists."""
        from app.services.protocols.template_converter import ConversionState

        conv_id = uuid4()
        state = ConversionState(test_org.id, conv_id)
        state.ensure_dir()
        fake_pdf = b"%PDF-1.4 fake preview content"
        state.write("preview.pdf", fake_pdf)

        resp = await client.get(
            f"/science/templates/conversions/{conv_id}/preview.pdf",
            headers=auth_headers,
        )
        assert resp.status_code == 200
        assert resp.content == fake_pdf
