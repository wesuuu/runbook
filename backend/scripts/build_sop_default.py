"""Generate the default SOP .docx template.

One-shot generator. Styling matches the time-course bioreactor SOP
example (light-gray KV header table, navy-blue section headings with a
horizontal rule underneath, bordered figure boxes, time-target tables);
the *content* is fully variable-driven — no example-specific copy.

Three rendering modes are supported on the same template via Jinja
conditionals:

  * ``is_time_based`` — time-course protocols render each section as
    a time-point block (e.g. ``3.1 Hour 0 (T=0)``) with an optional
    figure and a Time Target / Action / Expected Output table.
  * ``is_role_based`` — role-organized protocols fall through to the
    existing flat steps loop, one role per heading.
  * neither — a single flat steps table.

Run from ``backend/``:

    python scripts/build_sop_default.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUT = (
    Path(__file__).resolve().parent.parent
    / "app"
    / "services"
    / "documents"
    / "templates"
    / "sop_default.docx"
)

# Palette tuned to the bioreactor SOP example.
HEADER_FILL = "ECF0F1"      # cool light gray — KV / table headers
FIGURE_FILL = "FDFDFD"       # near-white — figure box interior
TEXT_BLACK = "000000"
TEXT_HEADING = "2C3E50"      # dark charcoal-navy — section headings
TEXT_MUTED = "5F6368"
TEXT_FIGURE_CAP = "5F6368"   # caption color
BORDER_GRAY = "BFBFBF"
RULE_GRAY = "D1D5DB"
FONT_NAME = "Arial"

# Vertical padding inside cells (twentieths of a point — w:tcMar units).
CELL_PAD_HEADER = 140        # ~7 pt top + bottom — table header rows
CELL_PAD_BODY = 100          # ~5 pt — KV + body rows


def _set_cell_shading(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _set_cell_border(cell, color: str = BORDER_GRAY, sz: int = 4) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(sz))
        b.set(qn("w:color"), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def _set_cell_margins(cell, *, top: int, bottom: int,
                      left: int = 100, right: int = 100) -> None:
    """Per-cell top/bottom/left/right margins in 1/20 pt (w:tcMar).

    Word ignores paragraph spacing inside table cells once they're laid
    out, so the only reliable way to "pad" a header row is via the cell
    margins element.
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:tcMar"))
    if existing is not None:
        tc_pr.remove(existing)
    tc_mar = OxmlElement("w:tcMar")
    for edge, val in (
        ("top", top), ("bottom", bottom),
        ("left", left), ("right", right),
    ):
        m = OxmlElement(f"w:{edge}")
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        tc_mar.append(m)
    tc_pr.append(tc_mar)


def _apply_font(run, *, size: int = 10, bold: bool = False,
                italic: bool = False, color: str = TEXT_BLACK):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    rfonts.set(qn("w:cs"), FONT_NAME)


def _style_cell(cell, text: str = "", *, bold: bool = False, size: int = 10,
                color: str = TEXT_BLACK, fill: str | None = None,
                align=None, pad: int = CELL_PAD_BODY):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    _apply_font(run, size=size, bold=bold, color=color)
    if fill:
        _set_cell_shading(cell, fill)
    _set_cell_border(cell)
    _set_cell_margins(cell, top=pad, bottom=pad)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _add_para(doc, text: str = "", *, bold: bool = False, italic: bool = False,
              size: int = 10, color: str = TEXT_BLACK, align=None,
              space_before: int | None = None,
              space_after: int | None = None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        _apply_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def _add_section_heading(doc, text: str, *, level: int = 2,
                         page_break_before: bool = False):
    """Navy heading with a horizontal rule underneath (set as a bottom
    border on the paragraph)."""
    size = {1: 22, 2: 16, 3: 13}.get(level, 13)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 2 else 12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    if page_break_before:
        br_run = p.add_run()
        br_run.add_break(WD_BREAK.PAGE)
    run = p.add_run(text)
    _apply_font(run, size=size, bold=True, color=TEXT_HEADING)

    # Bottom border = horizontal rule under the heading
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "8")
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), RULE_GRAY)
    pbdr.append(bot)
    p_pr.append(pbdr)
    return p


def _add_subsection_heading(doc, text: str):
    """Smaller navy heading for ``3.x`` time-point / role blocks. No
    underline rule — keeps the page lighter."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    _apply_font(run, size=13, bold=True, color=TEXT_HEADING)
    return p


def _add_jinja_para(doc, jinja: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(jinja)
    _apply_font(run, size=8, color=TEXT_MUTED)
    return p


def _set_table_layout(table, col_widths_pt: list[int] | None = None):
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    if col_widths_pt:
        for row in table.rows:
            for ci, w in enumerate(col_widths_pt):
                if ci < len(row.cells):
                    row.cells[ci].width = Pt(w)


def _figure_block(doc, caption_var: str, desc_var: str,
                  image_var: str | None = None):
    """Boxed figure: optional inline image, bold caption, italic
    description. Mirrors the bordered ``[Mock Figure N]`` blocks in the
    source SOP."""
    t = doc.add_table(rows=1, cols=1)
    _set_table_layout(t, [490])
    cell = t.rows[0].cells[0]
    cell.text = ""
    _set_cell_shading(cell, FIGURE_FILL)
    _set_cell_border(cell, color=BORDER_GRAY)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if image_var:
        img_p = cell.paragraphs[0]
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_p.paragraph_format.space_before = Pt(8)
        img_p.paragraph_format.space_after = Pt(4)
        _apply_font(img_p.add_run(image_var), size=10)
        cap_p = cell.add_paragraph()
    else:
        cap_p = cell.paragraphs[0]
        cap_p.paragraph_format.space_before = Pt(8)

    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_after = Pt(4)
    _apply_font(cap_p.add_run(caption_var),
                size=11, bold=True, color=TEXT_FIGURE_CAP)

    desc_p = cell.add_paragraph()
    desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc_p.paragraph_format.space_after = Pt(8)
    _apply_font(desc_p.add_run(desc_var),
                size=10, italic=True, color=TEXT_FIGURE_CAP)


def build() -> Document:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(TEXT_BLACK)

    for section in doc.sections:
        section.top_margin = Pt(54)
        section.bottom_margin = Pt(54)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

    # ── Unapproved warning banner ────────────────────────────────
    _add_jinja_para(doc, "{%p if unapproved_warning %}")
    warn = doc.add_paragraph()
    warn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_font(warn.add_run("⚠ UNAPPROVED — DRAFT ONLY"),
                size=11, bold=True, color="B91C1C")
    _add_jinja_para(doc, "{%p endif %}")

    # ── Title strip ──────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(6)
    _apply_font(title_p.add_run("Standard Operating Procedure"),
                size=12, bold=False, color=TEXT_MUTED)

    # ── Header KV table: doc # / revision / title / effective date ──
    th = doc.add_table(rows=3, cols=4)
    _set_table_layout(th, [120, 140, 100, 130])

    _style_cell(th.rows[0].cells[0], "Document Number:",
                bold=True, fill=HEADER_FILL)
    _style_cell(th.rows[0].cells[1], "{{ doc_number }}")
    _style_cell(th.rows[0].cells[2], "Revision:",
                bold=True, fill=HEADER_FILL)
    _style_cell(th.rows[0].cells[3], "{{ version_number }}")

    # Title row — merge value across 3 columns
    _style_cell(th.rows[1].cells[0], "Title:",
                bold=True, fill=HEADER_FILL)
    th.rows[1].cells[1].merge(th.rows[1].cells[2]).merge(th.rows[1].cells[3])
    _style_cell(th.rows[1].cells[1], "{{ protocol_name }}")

    # Effective date row — merge value across 3 columns
    _style_cell(th.rows[2].cells[0], "Effective Date:",
                bold=True, fill=HEADER_FILL)
    th.rows[2].cells[1].merge(th.rows[2].cells[2]).merge(th.rows[2].cells[3])
    _style_cell(th.rows[2].cells[1], "{{ effective_date }}")

    # Sub-meta line (project / org / generated) — small + muted
    _add_para(doc, "", space_after=4)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(8)
    _apply_font(
        meta.add_run(
            "{{ organization_name }}    "
            "Project: {{ project_name }}    "
            "Generated: {{ created_at }}"
        ),
        size=9, color=TEXT_MUTED,
    )

    # ── 1.0 Purpose ─────────────────────────────────────────────
    _add_section_heading(doc, "1.0 Purpose", level=2)
    purpose_p = doc.add_paragraph()
    purpose_p.paragraph_format.space_after = Pt(8)
    _apply_font(purpose_p.add_run("{{ purpose }}"), size=10)

    # ── 2.0 Scope ───────────────────────────────────────────────
    _add_section_heading(doc, "2.0 Scope", level=2)
    scope_p = doc.add_paragraph()
    scope_p.paragraph_format.space_after = Pt(8)
    _apply_font(scope_p.add_run("{{ scope }}"), size=10)

    # ── 3.0 Procedure ───────────────────────────────────────────
    _add_section_heading(doc, "3.0 Procedure: {{ protocol_name }}", level=2)

    # Optional critical-requirement callout
    _add_jinja_para(doc, "{%p if critical_requirement %}")
    crit = doc.add_paragraph()
    crit.paragraph_format.space_before = Pt(4)
    crit.paragraph_format.space_after = Pt(10)
    _apply_font(crit.add_run("CRITICAL REQUIREMENT: "),
                size=10, bold=True, color="B91C1C")
    _apply_font(crit.add_run("{{ critical_requirement }}"), size=10)
    _add_jinja_para(doc, "{%p endif %}")

    # ── 3.x time-course branch ──────────────────────────────────
    # When ``is_time_based`` is true, render each time point as a
    # subsection with optional figure + Time Target / Action / Output
    # table. Otherwise fall through to the role / flat-step branches.
    _add_jinja_para(doc, "{%p if is_time_based %}")
    _add_jinja_para(doc, "{%p for tp in time_points %}")

    _add_subsection_heading(
        doc,
        "3.{{ loop.index }}  {{ tp.name }}",
    )
    _add_jinja_para(doc, "{%p if tp.preamble %}")
    pre_p = doc.add_paragraph()
    pre_p.paragraph_format.space_after = Pt(6)
    _apply_font(pre_p.add_run("{{ tp.preamble }}"), size=10)
    _add_jinja_para(doc, "{%p endif %}")

    # Optional per-time-point figure
    _add_jinja_para(doc, "{%p if tp.figure %}")
    _figure_block(doc,
                  caption_var="{{ tp.figure.caption }}",
                  desc_var="{{ tp.figure.description }}",
                  image_var="{{ tp.figure.image }}")
    _add_para(doc, "", space_after=4)
    _add_jinja_para(doc, "{%p endif %}")

    # Time Target | Action | Expected Output table. Rows:
    #   0 header | 1 {%tr for %} | 2 content | 3 {%tr endfor %}
    tp_tbl = doc.add_table(rows=4, cols=3)
    _set_table_layout(tp_tbl, [110, 240, 140])
    for ci, h in enumerate(["Time Target", "Action", "Expected Output / Log"]):
        _style_cell(tp_tbl.rows[0].cells[ci], h, bold=True,
                    fill=HEADER_FILL, align=WD_ALIGN_PARAGRAPH.CENTER,
                    pad=CELL_PAD_HEADER)
    _style_cell(tp_tbl.rows[1].cells[0],
                "{%tr for act in tp.actions %}", size=8,
                color=TEXT_MUTED)
    _style_cell(tp_tbl.rows[1].cells[1], "")
    _style_cell(tp_tbl.rows[1].cells[2], "")
    _style_cell(tp_tbl.rows[2].cells[0], "{{ act.time }}", bold=True,
                align=WD_ALIGN_PARAGRAPH.CENTER)
    _style_cell(tp_tbl.rows[2].cells[1], "{{ act.action }}")
    _style_cell(tp_tbl.rows[2].cells[2], "{{ act.output }}",
                align=WD_ALIGN_PARAGRAPH.CENTER)
    _style_cell(tp_tbl.rows[3].cells[0], "{%tr endfor %}", size=8,
                color=TEXT_MUTED)
    _style_cell(tp_tbl.rows[3].cells[1], "")
    _style_cell(tp_tbl.rows[3].cells[2], "")

    _add_jinja_para(doc, "{%p endfor %}")
    _add_jinja_para(doc, "{%p else %}")

    # ── 3.x role-based / flat fallback ──────────────────────────
    _add_jinja_para(doc, "{%p if is_role_based and roles|length > 1 %}")
    _add_jinja_para(doc, "{%p for role in roles %}")
    _add_subsection_heading(doc,
                            "3.{{ loop.index }}  "
                            "{{ role.process_name or role.name }}")
    role_tbl = doc.add_table(rows=4, cols=4)
    _set_table_layout(role_tbl, [40, 130, 240, 80])
    for ci, h in enumerate(["Step", "Name", "Instruction", "Duration"]):
        _style_cell(role_tbl.rows[0].cells[ci], h, bold=True,
                    fill=HEADER_FILL, align=WD_ALIGN_PARAGRAPH.CENTER,
                    pad=CELL_PAD_HEADER)
    _style_cell(role_tbl.rows[1].cells[0],
                "{%tr for step in role.steps %}", size=8,
                color=TEXT_MUTED)
    for i in range(1, 4):
        _style_cell(role_tbl.rows[1].cells[i], "")
    _style_cell(role_tbl.rows[2].cells[0], "{{ loop.index }}",
                bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _style_cell(role_tbl.rows[2].cells[1], "{{ step.name }}")
    _style_cell(role_tbl.rows[2].cells[2], "{{ step.description }}")
    _style_cell(role_tbl.rows[2].cells[3],
                "{{ step.duration_min }} min",
                align=WD_ALIGN_PARAGRAPH.CENTER)
    _style_cell(role_tbl.rows[3].cells[0], "{%tr endfor %}", size=8,
                color=TEXT_MUTED)
    for i in range(1, 4):
        _style_cell(role_tbl.rows[3].cells[i], "")
    _add_jinja_para(doc, "{%p endfor %}")

    _add_jinja_para(doc, "{%p else %}")

    flat_tbl = doc.add_table(rows=4, cols=4)
    _set_table_layout(flat_tbl, [40, 130, 240, 80])
    for ci, h in enumerate(["Step", "Name", "Instruction", "Duration"]):
        _style_cell(flat_tbl.rows[0].cells[ci], h, bold=True,
                    fill=HEADER_FILL, align=WD_ALIGN_PARAGRAPH.CENTER,
                    pad=CELL_PAD_HEADER)
    _style_cell(flat_tbl.rows[1].cells[0],
                "{%tr for step in steps %}", size=8, color=TEXT_MUTED)
    for i in range(1, 4):
        _style_cell(flat_tbl.rows[1].cells[i], "")
    _style_cell(flat_tbl.rows[2].cells[0], "{{ loop.index }}",
                bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _style_cell(flat_tbl.rows[2].cells[1], "{{ step.name }}")
    _style_cell(flat_tbl.rows[2].cells[2], "{{ step.description }}")
    _style_cell(flat_tbl.rows[2].cells[3],
                "{{ step.duration_min }} min",
                align=WD_ALIGN_PARAGRAPH.CENTER)
    _style_cell(flat_tbl.rows[3].cells[0], "{%tr endfor %}", size=8,
                color=TEXT_MUTED)
    for i in range(1, 4):
        _style_cell(flat_tbl.rows[3].cells[i], "")

    _add_jinja_para(doc, "{%p endif %}")
    _add_jinja_para(doc, "{%p endif %}")

    # ── Approval block (preserved from previous template) ──────
    _add_jinja_para(doc, "{%p if approval %}")
    _add_section_heading(doc, "Approval & Signatures", level=2)
    appr_p = doc.add_paragraph()
    appr_p.paragraph_format.space_after = Pt(4)
    _apply_font(appr_p.add_run(
        "Approver: {{ approval.approver_name }} "
        "<{{ approval.approver_email }}>"
    ), size=10)
    appr2 = doc.add_paragraph()
    appr2.paragraph_format.space_after = Pt(4)
    _apply_font(appr2.add_run(
        "Approved at: {{ approval.approved_at }}    "
        "Protocol version: {{ approval.protocol_version }}"
    ), size=10, color=TEXT_MUTED)
    _add_jinja_para(doc, "{%p if approval.signature_statement %}")
    stmt = doc.add_paragraph()
    _apply_font(stmt.add_run(
        "Statement: {{ approval.signature_statement }}"
    ), size=10, italic=True)
    _add_jinja_para(doc, "{%p endif %}")
    _add_jinja_para(doc, "{%p if approval.signature_image %}")
    sig = doc.add_paragraph()
    _apply_font(sig.add_run("Signature: "), size=10, bold=True)
    _apply_font(sig.add_run("{{ approval.signature_image }}"), size=10)
    _add_jinja_para(doc, "{%p endif %}")
    _add_jinja_para(doc, "{%p endif %}")

    _add_jinja_para(doc, "{%p if approval_history %}")
    _add_section_heading(doc, "Approval History", level=2)
    hist_tbl = doc.add_table(rows=4, cols=3)
    _set_table_layout(hist_tbl, [140, 130, 220])
    for ci, h in enumerate(["Timestamp", "Action", "Actor"]):
        _style_cell(hist_tbl.rows[0].cells[ci], h, bold=True,
                    fill=HEADER_FILL, pad=CELL_PAD_HEADER)
    _style_cell(hist_tbl.rows[1].cells[0],
                "{%tr for ev in approval_history %}", size=8,
                color=TEXT_MUTED)
    _style_cell(hist_tbl.rows[1].cells[1], "")
    _style_cell(hist_tbl.rows[1].cells[2], "")
    _style_cell(hist_tbl.rows[2].cells[0], "{{ ev.created_at }}")
    _style_cell(hist_tbl.rows[2].cells[1], "{{ ev.action }}")
    _style_cell(hist_tbl.rows[2].cells[2], "{{ ev.actor_name }}")
    _style_cell(hist_tbl.rows[3].cells[0], "{%tr endfor %}", size=8,
                color=TEXT_MUTED)
    _style_cell(hist_tbl.rows[3].cells[1], "")
    _style_cell(hist_tbl.rows[3].cells[2], "")
    _add_jinja_para(doc, "{%p endif %}")

    return doc


def main() -> None:
    doc = build()
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT} ({OUT.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
