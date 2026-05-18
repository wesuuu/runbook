"""Generate the default batch record .docx template.

One-shot generator. Styling matches the GLP Tox Batch Manufacturing
Record example (light-gray table headers, Arial body, black text); the
*content* is fully driven by jinja variables wired to the keys produced
by ``template_engine.build_context`` — no example-specific copy.

Conditional sub-section tables in Section 4 (Unit Operations) only
render when the protocol is role-based AND has more than one role —
otherwise a single flat steps table is rendered.

Run from ``backend/``:

    python scripts/build_batch_record_default.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUT = (
    Path(__file__).resolve().parent.parent
    / "app"
    / "services"
    / "documents"
    / "templates"
    / "batch_record_default.docx"
)

# Match the GLP example's palette.
HEADER_FILL = "E8EAED"  # light gray — table header row
KV_FILL = "F1F3F4"  # very light gray — KV label cells
TEXT_BLACK = "000000"
TEXT_MUTED = "5F6368"
BORDER_GRAY = "BFBFBF"
FONT_NAME = "Arial"


def _set_cell_shading(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _set_cell_border(cell, color: str = BORDER_GRAY, sz: int = 4) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(sz))
        b.set(qn("w:color"), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def _apply_font(
    run,
    *,
    size: int = 10,
    bold: bool = False,
    italic: bool = False,
    color: str = TEXT_BLACK,
):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    # Force the font for east-asian / complex script too
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    rfonts.set(qn("w:cs"), FONT_NAME)


def _style_cell(
    cell,
    text: str = "",
    *,
    bold: bool = False,
    size: int = 10,
    color: str = TEXT_BLACK,
    fill: str | None = None,
    align=None,
):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    _apply_font(run, size=size, bold=bold, color=color)
    if fill:
        _set_cell_shading(cell, fill)
    _set_cell_border(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _add_para(
    doc,
    text: str = "",
    *,
    bold: bool = False,
    italic: bool = False,
    size: int = 10,
    color: str = TEXT_BLACK,
    align=None,
    space_before: int | None = None,
    space_after: int | None = None,
):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        _apply_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def _add_heading(doc, text: str, *, level: int = 2, page_break_before: bool = False):
    """Section heading. ``page_break_before=True`` puts the heading on a
    new page using a w:br w:type='page' run, which both Word and
    LibreOffice render reliably."""
    size = {1: 16, 2: 13, 3: 11}.get(level, 11)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 2 else 10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    if page_break_before:
        br_run = p.add_run()
        br_run.add_break(WD_BREAK.PAGE)
    run = p.add_run(text)
    _apply_font(run, size=size, bold=True, color=TEXT_BLACK)
    return p


def _add_jinja_para(doc, jinja: str):
    """Carrier paragraph for a docxtpl control tag (e.g.
    ``{%p for role in roles %}``). docxtpl strips the carrier."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(jinja)
    _apply_font(run, size=8, color=TEXT_MUTED)
    return p


def _set_table_layout(table, col_widths_pt: list[int] | None = None):
    """Fix table to defined column widths; helps LibreOffice + Word
    render with the same proportions."""
    # Disable auto-fit
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    if col_widths_pt:
        for row in table.rows:
            for ci, w in enumerate(col_widths_pt):
                if ci < len(row.cells):
                    row.cells[ci].width = Pt(w)


def build() -> Document:
    doc = Document()

    # Style the default Normal style — applies to anything not
    # explicitly overridden (e.g., empty carrier paragraphs).
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(TEXT_BLACK)

    # Page margins (US Letter, ~0.75" / 0.5" sides)
    for section in doc.sections:
        section.top_margin = Pt(54)
        section.bottom_margin = Pt(54)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

    # ── Unapproved warning banner ────────────────────────────────
    _add_jinja_para(doc, "{%p if unapproved_warning %}")
    warn = doc.add_paragraph()
    warn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    wrun = warn.add_run("⚠ UNAPPROVED — DRAFT ONLY")
    _apply_font(wrun, size=11, bold=True, color="B91C1C")
    _add_jinja_para(doc, "{%p endif %}")

    # ── Document title (all variable-driven) ─────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(2)
    _apply_font(
        title_p.add_run("{{ protocol_name }}"), size=18, bold=True, color=TEXT_BLACK
    )

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(2)
    _apply_font(
        sub_p.add_run("Batch Manufacturing Record"),
        size=11,
        italic=True,
        color=TEXT_MUTED,
    )

    org_p = doc.add_paragraph()
    org_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org_p.paragraph_format.space_after = Pt(8)
    _apply_font(org_p.add_run("{{ organization_name }}"), size=10, color=TEXT_MUTED)

    # Meta line(s) — run / project / version / dates. All variable.
    _add_jinja_para(doc, "{%p if run_name %}")
    meta_run = doc.add_paragraph()
    meta_run.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_font(
        meta_run.add_run(
            "Run: {{ run_name }}    Status: {{ run_status }}    "
            "Started: {{ started_at }}    Completed: {{ completed_at }}"
        ),
        size=9,
        color=TEXT_MUTED,
    )
    _add_jinja_para(doc, "{%p endif %}")

    meta_proj = doc.add_paragraph()
    meta_proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_proj.paragraph_format.space_after = Pt(6)
    _apply_font(
        meta_proj.add_run(
            "Project: {{ project_name }}    "
            "Protocol Version: v{{ version_number }}    "
            "Generated: {{ created_at }}"
        ),
        size=9,
        color=TEXT_MUTED,
    )

    # ── Section 1: General Information ───────────────────────────
    _add_heading(doc, "1. General Information", level=2)

    t1 = doc.add_table(rows=4, cols=2)
    _set_table_layout(t1, [170, 320])
    # F-0087: target_yield removed from the header per GLP plan — the
    # run outcome is the relevant disposition once the batch is complete.
    rows_data = [
        ("Product Name / Candidate", "{{ protocol_name }}"),
        ("Batch / Lot Number", "{{ run_name }}"),
        ("Outcome", "{{ run.outcome }}"),
        (
            "Process SOP / Protocol Reference",
            "{{ protocol_name }} (v{{ version_number }})",
        ),
    ]
    for r, (label, value) in enumerate(rows_data):
        _style_cell(
            t1.rows[r].cells[0],
            label,
            bold=True,
            size=10,
            color=TEXT_BLACK,
            fill=KV_FILL,
        )
        _style_cell(t1.rows[r].cells[1], value, size=10)

    # F-0087: outcome notes paragraph (conditional, below the header).
    _add_jinja_para(doc, "{%p if run.outcome_notes %}")
    notes_label = doc.add_paragraph()
    notes_label.paragraph_format.space_before = Pt(6)
    notes_label.paragraph_format.space_after = Pt(2)
    _apply_font(
        notes_label.add_run("Outcome Notes:"), size=10, bold=True, color=TEXT_BLACK
    )
    notes_p = doc.add_paragraph()
    notes_p.paragraph_format.space_after = Pt(6)
    _apply_font(notes_p.add_run("{{ run.outcome_notes }}"), size=10, color=TEXT_BLACK)
    _add_jinja_para(doc, "{%p endif %}")

    # ── Section 2: Bill of Materials ─────────────────────────────
    _add_heading(doc, "2. Bill of Materials (BOM)", level=2)
    _add_para(
        doc,
        "Record all raw materials, media, and buffers utilized in this run.",
        size=9,
        color=TEXT_MUTED,
        italic=True,
        space_after=4,
    )

    t2 = doc.add_table(rows=5, cols=5)
    _set_table_layout(t2, [150, 100, 80, 80, 90])
    bom_headers = [
        "Material Description",
        "Part / Lot Number",
        "Required Quantity",
        "Actual Quantity Added",
        "Operator Initials & Date",
    ]
    for ci, h in enumerate(bom_headers):
        _style_cell(
            t2.rows[0].cells[ci],
            h,
            bold=True,
            size=10,
            color=TEXT_BLACK,
            fill=HEADER_FILL,
        )
    _style_cell(
        t2.rows[1].cells[0],
        "{%tr for material in materials %}",
        size=8,
        color=TEXT_MUTED,
    )
    _style_cell(t2.rows[2].cells[0], "{{ material.description }}", size=10)
    _style_cell(t2.rows[2].cells[1], "{{ material.lot_number }}", size=10)
    _style_cell(t2.rows[2].cells[2], "{{ material.required_qty }}", size=10)
    _style_cell(t2.rows[2].cells[3], "{{ material.actual_qty }}", size=10)
    _style_cell(t2.rows[2].cells[4], "{{ material.operator }}", size=10)
    _style_cell(t2.rows[3].cells[0], "{%tr endfor %}", size=8, color=TEXT_MUTED)
    for ci in range(5):
        _style_cell(t2.rows[4].cells[ci], "", size=10)

    # ── Section 3: Equipment Log ─────────────────────────────────
    _add_heading(doc, "3. Equipment Log", level=2)
    _add_para(
        doc,
        "Record primary equipment used to ensure traceability.",
        size=9,
        color=TEXT_MUTED,
        italic=True,
        space_after=4,
    )

    t3 = doc.add_table(rows=5, cols=4)
    _set_table_layout(t3, [160, 110, 130, 100])
    eq_headers = [
        "Equipment Name / Type",
        "Equipment ID (Asset Tag)",
        "Calibration Status (Valid Until)",
        "Operator Initials & Date",
    ]
    for ci, h in enumerate(eq_headers):
        _style_cell(
            t3.rows[0].cells[ci],
            h,
            bold=True,
            size=10,
            color=TEXT_BLACK,
            fill=HEADER_FILL,
        )
    _style_cell(
        t3.rows[1].cells[0], "{%tr for eq in equipment %}", size=8, color=TEXT_MUTED
    )
    _style_cell(t3.rows[2].cells[0], "{{ eq.name }}", size=10)
    _style_cell(t3.rows[2].cells[1], "{{ eq.asset_id }}", size=10)
    _style_cell(t3.rows[2].cells[2], "{{ eq.calibration }}", size=10)
    _style_cell(t3.rows[2].cells[3], "{{ eq.operator }}", size=10)
    _style_cell(t3.rows[3].cells[0], "{%tr endfor %}", size=8, color=TEXT_MUTED)
    for ci in range(4):
        _style_cell(t3.rows[4].cells[ci], "", size=10)

    # ── Section 4: Execution: Unit Operations ───────────────────
    _add_heading(doc, "4. Execution: Unit Operations", level=2)
    _add_para(
        doc,
        "Execute process steps. Record actual values, operator, and "
        "verifier (if required by SOP) for each step.",
        size=9,
        color=TEXT_MUTED,
        italic=True,
        space_after=4,
    )

    step_headers = [
        "Step",
        "Instruction",
        "Target Parameter",
        "Actual Value",
        "Operator",
        "Verifier",
    ]
    step_widths = [40, 130, 130, 80, 50, 50]

    def _emit_step_table(loop_expr: str):
        t = doc.add_table(rows=4, cols=6)
        _set_table_layout(t, step_widths)
        for ci, h in enumerate(step_headers):
            _style_cell(
                t.rows[0].cells[ci],
                h,
                bold=True,
                size=10,
                color=TEXT_BLACK,
                fill=HEADER_FILL,
            )
        _style_cell(t.rows[1].cells[0], loop_expr, size=8, color=TEXT_MUTED)
        _style_cell(
            t.rows[2].cells[0],
            "{{ loop.index }}",
            size=10,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _style_cell(t.rows[2].cells[1], "{{ step.name }}", size=10)
        _style_cell(t.rows[2].cells[2], "{{ step.description }}", size=9)
        # F-0087: GLP-compliant recorded value. ``actual_value_block``
        # is a multi-line plain string (target / recorded / operator •
        # timestamp). ``edit_reason`` surfaces the GMP-style rationale
        # when the value was edited post-completion.
        cell_val = t.rows[2].cells[3]
        cell_val.text = ""
        p_val = cell_val.paragraphs[0]
        run_val = p_val.add_run("{{r step.actual_value_block }}")
        _apply_font(run_val, size=10)
        # Edit-reason annotation in a second paragraph (conditional via
        # docxtpl ``{%p if %}`` so an unedited step renders no extra row).
        er_open_p = cell_val.add_paragraph()
        er_open_p.paragraph_format.space_before = Pt(0)
        er_open_p.paragraph_format.space_after = Pt(0)
        _apply_font(
            er_open_p.add_run("{%p if step.edit_reason %}"),
            size=8,
            color=TEXT_MUTED,
        )
        er_body = cell_val.add_paragraph()
        er_body.paragraph_format.space_before = Pt(0)
        er_body.paragraph_format.space_after = Pt(0)
        _apply_font(
            er_body.add_run("Edit reason: "), size=8, bold=True, color=TEXT_MUTED
        )
        _apply_font(
            er_body.add_run("{{ step.edit_reason }}"),
            size=8,
            italic=True,
            color=TEXT_MUTED,
        )
        er_close_p = cell_val.add_paragraph()
        er_close_p.paragraph_format.space_before = Pt(0)
        er_close_p.paragraph_format.space_after = Pt(0)
        _apply_font(
            er_close_p.add_run("{%p endif %}"),
            size=8,
            color=TEXT_MUTED,
        )
        _set_cell_border(cell_val)
        cell_val.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _style_cell(
            t.rows[2].cells[4],
            "{{ step.initials }}",
            size=10,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _style_cell(t.rows[2].cells[5], "", size=10)
        _style_cell(t.rows[3].cells[0], "{%tr endfor %}", size=8, color=TEXT_MUTED)

    _add_jinja_para(doc, "{%p if is_role_based and roles|length > 1 %}")
    _add_jinja_para(doc, "{%p for role in roles %}")
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(8)
    sub_p.paragraph_format.space_after = Pt(4)
    sub_p.paragraph_format.keep_with_next = True
    _apply_font(
        sub_p.add_run("4.{{ loop.index }}  {{ role.process_name or role.name }}"),
        size=11,
        bold=True,
        color=TEXT_BLACK,
    )
    _emit_step_table("{%tr for step in role.steps %}")
    _add_jinja_para(doc, "{%p endfor %}")
    _add_jinja_para(doc, "{%p else %}")
    _emit_step_table("{%tr for step in steps %}")
    _add_jinja_para(doc, "{%p endif %}")

    # ── Section 5: Deviations and Process Comments ──────────────
    _add_heading(doc, "5. Deviations and Process Comments", level=2)
    _add_para(
        doc,
        "Log any deviations from target parameters or SOP. Run notes "
        "and anomalies are surfaced here.",
        size=9,
        color=TEXT_MUTED,
        italic=True,
        space_after=4,
    )

    t5 = doc.add_table(rows=5, cols=4)
    _set_table_layout(t5, [90, 220, 90, 100])
    dev_headers = [
        "Step / Status Ref.",
        "Description of Deviation / Observation",
        "Impact Assessment Required? (Y/N)",
        "Lead Reviewer Sign-off",
    ]
    for ci, h in enumerate(dev_headers):
        _style_cell(
            t5.rows[0].cells[ci],
            h,
            bold=True,
            size=10,
            color=TEXT_BLACK,
            fill=HEADER_FILL,
        )
    _style_cell(
        t5.rows[1].cells[0], "{%tr for note in notes %}", size=8, color=TEXT_MUTED
    )
    _style_cell(t5.rows[2].cells[0], "{{ note.run_status }}", size=10)
    _style_cell(t5.rows[2].cells[1], "{{ note.content }}", size=10)
    _style_cell(t5.rows[2].cells[2], "", size=10)
    _style_cell(
        t5.rows[2].cells[3], "{{ note.author_name }}\n{{ note.created_at }}", size=9
    )
    _style_cell(t5.rows[3].cells[0], "{%tr endfor %}", size=8, color=TEXT_MUTED)
    for ci in range(4):
        _style_cell(t5.rows[4].cells[ci], "", size=10)

    # ── Section 6: Final Disposition & Signatures ───────────────
    _add_heading(doc, "6. Final Disposition & Signatures", level=2)
    _add_para(
        doc,
        "By signing below, the responsible parties certify that the "
        "material described herein was produced according to the specified "
        "process and that all deviations have been reviewed.",
        size=9,
        color=TEXT_MUTED,
        italic=True,
        space_after=6,
    )

    # ── Protocol Approvals (F-0087 per-role) ────────────────────
    # One block per role from ``protocol_approvals.<role>``.
    _add_para(
        doc, "Protocol Approvals", bold=True, size=11, space_before=4, space_after=4
    )

    def _emit_approval_role(role_key: str, label: str, cfr_cite: str):
        _add_jinja_para(doc, f"{{%p if protocol_approvals.{role_key} %}}")
        hdr = doc.add_paragraph()
        hdr.paragraph_format.space_before = Pt(6)
        hdr.paragraph_format.space_after = Pt(2)
        hdr.paragraph_format.keep_with_next = True
        _apply_font(hdr.add_run(f"{label} "), size=10, bold=True, color=TEXT_BLACK)
        _apply_font(hdr.add_run(f"({cfr_cite})"), size=9, color=TEXT_MUTED)
        name_p = doc.add_paragraph()
        name_p.paragraph_format.space_after = Pt(2)
        _apply_font(
            name_p.add_run(
                f"Name: {{{{ protocol_approvals.{role_key}.name }}}}    "
                f"Signed: "
                f"{{{{ protocol_approvals.{role_key}.signed_at }}}}"
            ),
            size=9,
            color=TEXT_BLACK,
        )
        _add_jinja_para(doc, f"{{%p if protocol_approvals.{role_key}.attestation %}}")
        att_p = doc.add_paragraph()
        att_p.paragraph_format.space_after = Pt(2)
        _apply_font(att_p.add_run("Attestation: "), size=9, bold=True)
        _apply_font(
            att_p.add_run(f"{{{{ protocol_approvals.{role_key}.attestation }}}}"),
            size=9,
            italic=True,
        )
        _add_jinja_para(doc, "{%p endif %}")
        _add_jinja_para(
            doc,
            f"{{%p if protocol_approvals.{role_key}.signature_image %}}",
        )
        sig_p = doc.add_paragraph()
        sig_p.paragraph_format.space_after = Pt(4)
        _apply_font(sig_p.add_run("Signature: "), size=9, bold=True)
        _apply_font(
            sig_p.add_run(f"{{{{ protocol_approvals.{role_key}.signature_image }}}}"),
            size=9,
        )
        _add_jinja_para(doc, "{%p endif %}")
        _add_jinja_para(doc, "{%p endif %}")

    _emit_approval_role("sponsor", "Sponsor", "21 CFR §58.10")
    _emit_approval_role("study_director", "Study Director", "21 CFR §58.33")
    _emit_approval_role("qau", "Quality Assurance Unit", "21 CFR §58.35")

    # ── Run Sign-offs (F-0087, new) ────────────────────────────
    # One block per GLP role from ``signoffs.<role>``: operator, study
    # director, QAU. Each block only renders if a sign-off exists.
    _add_para(doc, "Run Sign-offs", bold=True, size=11, space_before=10, space_after=4)

    def _emit_signoff_role(role_key: str, label: str, cfr_cite: str):
        _add_jinja_para(doc, f"{{%p if signoffs.{role_key} %}}")
        hdr = doc.add_paragraph()
        hdr.paragraph_format.space_before = Pt(6)
        hdr.paragraph_format.space_after = Pt(2)
        hdr.paragraph_format.keep_with_next = True
        _apply_font(hdr.add_run(f"{label} "), size=10, bold=True, color=TEXT_BLACK)
        _apply_font(hdr.add_run(f"({cfr_cite})"), size=9, color=TEXT_MUTED)
        name_p = doc.add_paragraph()
        name_p.paragraph_format.space_after = Pt(2)
        _apply_font(
            name_p.add_run(
                f"Name: {{{{ signoffs.{role_key}.name }}}}    "
                f"Signed: {{{{ signoffs.{role_key}.signed_at }}}}"
            ),
            size=9,
            color=TEXT_BLACK,
        )
        _add_jinja_para(doc, f"{{%p if signoffs.{role_key}.attestation %}}")
        att_p = doc.add_paragraph()
        att_p.paragraph_format.space_after = Pt(2)
        _apply_font(att_p.add_run("Attestation: "), size=9, bold=True)
        _apply_font(
            att_p.add_run(f"{{{{ signoffs.{role_key}.attestation }}}}"),
            size=9,
            italic=True,
        )
        _add_jinja_para(doc, "{%p endif %}")
        _add_jinja_para(doc, f"{{%p if signoffs.{role_key}.signature_image %}}")
        sig_p = doc.add_paragraph()
        sig_p.paragraph_format.space_after = Pt(4)
        _apply_font(sig_p.add_run("Signature: "), size=9, bold=True)
        _apply_font(
            sig_p.add_run(f"{{{{ signoffs.{role_key}.signature_image }}}}"),
            size=9,
        )
        _add_jinja_para(doc, "{%p endif %}")
        _add_jinja_para(doc, "{%p endif %}")

    _emit_signoff_role("operator", "Operator", "21 CFR §58.29")
    _emit_signoff_role("study_director", "Study Director", "21 CFR §58.33")
    _emit_signoff_role("qau", "Quality Assurance Unit", "21 CFR §58.35")

    # Manual wet-ink sign-off rows — always present (mirrors GLP)
    _add_para(
        doc, "Wet-Ink Sign-Off", bold=True, size=11, space_before=8, space_after=4
    )
    t7 = doc.add_table(rows=4, cols=4)
    _set_table_layout(t7, [150, 130, 130, 90])
    wet_headers = ["Role", "Name (Print)", "Signature", "Date"]
    for ci, h in enumerate(wet_headers):
        _style_cell(
            t7.rows[0].cells[ci],
            h,
            bold=True,
            size=10,
            color=TEXT_BLACK,
            fill=HEADER_FILL,
        )
    wet_roles = [
        "Lead Operator / Engineer",
        "Process Development Lead",
        "Quality / Compliance (if applicable)",
    ]
    for ri, role in enumerate(wet_roles, start=1):
        _style_cell(t7.rows[ri].cells[0], role, bold=True, size=10, fill=KV_FILL)
        for ci in range(1, 4):
            _style_cell(t7.rows[ri].cells[ci], "", size=10)

    # ── Appendix A: Figures ─────────────────────────────────────
    _add_jinja_para(doc, "{%p if figures %}")
    _add_heading(doc, "Appendix A: Figures", level=2, page_break_before=True)
    _add_jinja_para(doc, "{%p for fig in figures %}")
    fig_p = doc.add_paragraph()
    fig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_p.paragraph_format.space_before = Pt(6)
    fig_p.paragraph_format.space_after = Pt(2)
    fig_p.paragraph_format.keep_with_next = True
    _apply_font(fig_p.add_run("{{ fig.image }}"), size=10)
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_after = Pt(8)
    _apply_font(
        cap_p.add_run(
            "Figure {{ fig.number }}: {{ fig.filename }}  "
            "({{ fig.step_name }}, uploaded {{ fig.uploaded_at }})"
        ),
        size=9,
        italic=True,
        color=TEXT_MUTED,
    )
    _add_jinja_para(doc, "{%p endfor %}")
    _add_jinja_para(doc, "{%p endif %}")

    # ── Appendix B: Non-image Attachments ───────────────────────
    _add_jinja_para(doc, "{%p if non_image_attachments %}")
    _add_heading(doc, "Appendix B: Attachments", level=2, page_break_before=True)
    t8 = doc.add_table(rows=3, cols=4)
    _set_table_layout(t8, [180, 90, 130, 100])
    att_headers = ["Filename", "Type", "Scope", "Uploaded"]
    for ci, h in enumerate(att_headers):
        _style_cell(
            t8.rows[0].cells[ci],
            h,
            bold=True,
            size=10,
            color=TEXT_BLACK,
            fill=HEADER_FILL,
        )
    _style_cell(
        t8.rows[1].cells[0],
        "{%tr for att in non_image_attachments %}",
        size=8,
        color=TEXT_MUTED,
    )
    _style_cell(t8.rows[2].cells[0], "{{ att.filename }}", size=10)
    _style_cell(t8.rows[2].cells[1], "{{ att.type }}", size=10)
    _style_cell(t8.rows[2].cells[2], "{{ att.scope }}", size=10)
    _style_cell(t8.rows[2].cells[3], "{{ att.uploaded_at }}", size=10)
    t8.add_row()
    _style_cell(t8.rows[3].cells[0], "{%tr endfor %}", size=8, color=TEXT_MUTED)
    _add_jinja_para(doc, "{%p endif %}")

    return doc


def main() -> None:
    doc = build()
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Wrote {OUT} ({OUT.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
