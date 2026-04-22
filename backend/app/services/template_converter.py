"""Template conversion agent — converts filled DOCX documents to Jinja2 templates.

Uses a substitution-based approach: the AI identifies variable values in the
original document and returns find-and-replace pairs. The code performs the
substitutions on the DOCX directly, preserving all original formatting.
Progress is streamed to the frontend via Server-Sent Events (SSE).
"""

import asyncio
import json
import logging
import re
import subprocess
import tempfile
import time
from collections.abc import AsyncIterator
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Any
from uuid import UUID, uuid4

from ulid import ULID

from docxtpl import DocxTemplate
from pydantic_ai import Agent, RunContext
from pydantic_ai.messages import BinaryContent
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.services.file_storage import FileStorageService
from app.services.template_engine import KNOWN_VARIABLES, get_mock_context

logger = logging.getLogger(__name__)

JINJA_PATTERN = re.compile(r"\{\{.*?\}\}|\{%.*?%\}")

TOOL_LIMIT_MSG = (
    "Tool call limit ({limit}) reached. Conversion cannot continue — "
    "the document may be too complex for the current model. "
    "Please simplify the document or configure a more capable AI model."
)


# ── SSE Event Stream ──


class EventStream:
    """Buffered async event stream for SSE.

    The background conversion task pushes events via push().
    The SSE endpoint reads events via the async iterator protocol.
    Supports late-joining clients by replaying from buffer index 0.
    """

    def __init__(self) -> None:
        self._events: list[tuple[str, str]] = []
        self._waiters: list[asyncio.Event] = []
        self._closed = False

    def push(self, event_type: str, data: dict) -> None:
        """Push an event. Called from the background task."""
        self._events.append((event_type, json.dumps(data, default=str)))
        for w in self._waiters:
            w.set()

    def close(self) -> None:
        """Signal no more events will come."""
        self._closed = True
        for w in self._waiters:
            w.set()

    async def iter_events(self) -> AsyncIterator[tuple[str, str]]:
        """Async iterator yielding (event_type, data_json) tuples."""
        idx = 0
        while True:
            while idx < len(self._events):
                yield self._events[idx]
                idx += 1
            if self._closed:
                return
            waiter = asyncio.Event()
            self._waiters.append(waiter)
            await waiter.wait()
            self._waiters.remove(waiter)


# Module-level registry: conversion_id -> EventStream
_active_streams: dict[str, EventStream] = {}


# ── Data Classes ──


@dataclass
class RenderResult:
    """Result of attempting to render a template with mock data."""

    pdf_bytes: bytes
    jinja_remnants: list[str]
    format_issues: list[str]
    render_error: str | None = None


@dataclass
class ConversionDeps:
    """Dependencies injected into pydantic-ai tools via RunContext."""

    state: "ConversionState"
    event_stream: EventStream
    org_id: UUID
    template_type: str
    original_pdf_bytes: bytes
    model: Any
    tool_call_count: int = field(default=0, init=False)
    max_tool_calls: int = field(
        default_factory=lambda: settings.template_convert_max_tool_calls
    )
    trace_entries: list[dict] = field(default_factory=list, init=False)
    started_at: float = field(default_factory=time.time, init=False)


class ConversionState:
    """Tracks state for an active conversion session on the filesystem.

    Files are stored under:
        {storage_root}/{org_id}/tmp/conversions/{conversion_id}/
    """

    def __init__(
        self,
        org_id: UUID,
        conversion_id: UUID | str,
        storage_root: str = "./uploads",
    ) -> None:
        self.org_id = org_id
        self.conversion_id = conversion_id
        self.storage_root = Path(storage_root)
        self.base_path = (
            Path(str(org_id)) / "tmp" / "conversions" / str(conversion_id)
        )

    def _resolve(self, filename: str) -> Path:
        """Resolve a filename to full path under the conversion directory."""
        full = (self.storage_root / self.base_path / filename).resolve()
        root = self.storage_root.resolve()
        if not str(full).startswith(str(root) + "/") and full != root:
            raise ValueError("Path traversal detected")
        return full

    def ensure_dir(self) -> None:
        """Create the conversion directory if it doesn't exist."""
        self._resolve("_").parent.mkdir(parents=True, exist_ok=True)

    def write(self, filename: str, data: bytes) -> None:
        self._resolve(filename).write_bytes(data)

    def read(self, filename: str) -> bytes:
        return self._resolve(filename).read_bytes()

    def exists(self, filename: str) -> bool:
        return self._resolve(filename).exists()

    def write_json(self, filename: str, data: dict | list) -> None:
        self._resolve(filename).write_text(
            json.dumps(data, default=str), encoding="utf-8"
        )

    def read_json(self, filename: str) -> dict | list:
        return json.loads(self._resolve(filename).read_text(encoding="utf-8"))

    @property
    def preview_url(self) -> str:
        return (
            f"/science/templates/conversions/"
            f"{self.conversion_id}/preview.pdf"
        )

    @property
    def template_url(self) -> str:
        return (
            f"/science/templates/conversions/"
            f"{self.conversion_id}/template.docx"
        )


# ── Helper Functions ──


def _extract_jinja_variables(text: str) -> set[str]:
    """Extract Jinja2 variable names from template text."""
    var_pattern = re.compile(r"\{\{\s*([\w.]+)\s*\}\}")
    loop_pattern = re.compile(
        r"\{%\s*(?:tr\s+)?for\s+(\w+)\s+in\s+(\w+)\s*%\}"
    )
    variables: set[str] = set()

    for match in var_pattern.finditer(text):
        name = match.group(1)
        top = name.split(".")[0]
        variables.add(top)

    for match in loop_pattern.finditer(text):
        variables.add(match.group(1))
        variables.add(match.group(2))

    return variables


def _try_render(template_bytes: bytes, template_type: str) -> RenderResult:
    """Render a template DOCX with mock data and check for issues."""
    mock_ctx = get_mock_context()

    # Add placeholder values for template variables not in mock_ctx.
    # Also replace complex mock values used directly (not in loops)
    # with readable placeholders.
    all_text = _extract_docx_text(template_bytes)
    detected = _extract_jinja_variables(all_text)

    loop_collections: set[str] = set()
    loop_pat = re.compile(
        r"\{%\s*(?:tr\s+)?for\s+\w+\s+in\s+(\w+)\s*%\}"
    )
    for match in loop_pat.finditer(all_text):
        loop_collections.add(match.group(1))

    for var in detected:
        if var not in mock_ctx:
            mock_ctx[var] = f"[{var}]"
        elif var not in loop_collections and isinstance(
            mock_ctx[var], (list, dict)
        ):
            mock_ctx[var] = f"[{var}]"

    with tempfile.TemporaryDirectory() as tmpdir:
        tpl_path = Path(tmpdir) / "template.docx"
        tpl_path.write_bytes(template_bytes)

        try:
            doc = DocxTemplate(str(tpl_path))
            doc.render(mock_ctx)
            rendered_path = Path(tmpdir) / "rendered.docx"
            doc.save(str(rendered_path))
        except Exception as e:
            return RenderResult(
                pdf_bytes=b"",
                jinja_remnants=[],
                format_issues=[],
                render_error=f"docxtpl render failed: {e}",
            )

        from docx import Document

        rendered_doc = Document(str(rendered_path))
        all_text_parts: list[str] = []
        for p in rendered_doc.paragraphs:
            all_text_parts.append(p.text)
        for table in rendered_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text_parts.append(cell.text)

        rendered_text = "\n".join(all_text_parts)
        remnants = JINJA_PATTERN.findall(rendered_text)

        pdf_bytes = b""
        try:
            subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    tmpdir,
                    str(rendered_path),
                ],
                capture_output=True,
                timeout=30,
            )
            pdf_path = Path(tmpdir) / "rendered.pdf"
            if pdf_path.exists():
                pdf_bytes = pdf_path.read_bytes()
        except (subprocess.TimeoutExpired, FileNotFoundError) as e:
            logger.warning("PDF conversion failed: %s", e)

        return RenderResult(
            pdf_bytes=pdf_bytes,
            jinja_remnants=remnants,
            format_issues=[],
        )


def _extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract readable text content from a DOCX file."""
    from docx import Document as DocxDocument

    doc = DocxDocument(BytesIO(file_bytes))
    parts: list[str] = []

    for element in doc.element.body:
        tag = (
            element.tag.split("}")[-1]
            if "}" in element.tag
            else element.tag
        )
        if tag == "p":
            runs_text = ""
            for run in element.iter():
                run_tag = (
                    run.tag.split("}")[-1]
                    if "}" in run.tag
                    else run.tag
                )
                if run_tag == "t" and run.text:
                    runs_text += run.text
            if runs_text.strip():
                parts.append(runs_text.strip())
        elif tag == "tbl":
            parts.append("[TABLE]")
            for row in element.iter():
                row_tag = (
                    row.tag.split("}")[-1]
                    if "}" in row.tag
                    else row.tag
                )
                if row_tag == "tr":
                    cells: list[str] = []
                    for cell in row.iter():
                        cell_tag = (
                            cell.tag.split("}")[-1]
                            if "}" in cell.tag
                            else cell.tag
                        )
                        if cell_tag == "tc":
                            cell_text = ""
                            for t in cell.iter():
                                t_tag = (
                                    t.tag.split("}")[-1]
                                    if "}" in t.tag
                                    else t.tag
                                )
                                if t_tag == "t" and t.text:
                                    cell_text += t.text
                            cells.append(cell_text.strip())
                    if cells:
                        parts.append(" | ".join(cells))
            parts.append("[/TABLE]")

    return "\n".join(parts)


def _extract_docx_text(template_bytes: bytes) -> str:
    """Extract all text from a DOCX file (paragraphs + tables)."""
    from docx import Document as DocxDocument

    doc = DocxDocument(BytesIO(template_bytes))
    parts: list[str] = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


async def _to_pdf(file_bytes: bytes, filename: str) -> bytes:
    """Convert any supported file to PDF via LibreOffice."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return file_bytes
    if ext in (".png", ".jpg", ".jpeg"):
        return file_bytes

    with tempfile.TemporaryDirectory() as tmpdir:
        input_path = Path(tmpdir) / f"input{ext}"
        input_path.write_bytes(file_bytes)
        result = subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                tmpdir,
                str(input_path),
            ],
            capture_output=True,
            timeout=60,
        )
        pdf_path = input_path.with_suffix(".pdf")
        if not pdf_path.exists():
            raise ValueError(
                f"LibreOffice failed to convert {filename} to PDF: "
                f"{result.stderr.decode()}"
            )
        return pdf_path.read_bytes()


def _replace_in_paragraph(paragraph: Any, find: str, replace: str) -> bool:
    """Replace text that may span multiple runs in a paragraph.

    DOCX often splits text across runs (due to formatting, spell-check,
    edit history). This function handles cross-run replacements by:
    1. Trying single-run replacement first (preserves formatting best)
    2. Falling back to cross-run replacement (consolidates into first
       matching run, clears subsequent runs that were part of the match)

    Returns True if a replacement was made.
    """
    if find not in paragraph.text:
        return False

    # Fast path: text is within a single run
    for run in paragraph.runs:
        if find in run.text:
            run.text = run.text.replace(find, replace)
            return True

    # Slow path: text spans multiple runs — find which runs contain it
    runs = paragraph.runs
    if not runs:
        return False

    # Build a map of character positions to run indices
    full_text = ""
    run_boundaries: list[tuple[int, int, int]] = []  # (start, end, idx)
    for i, run in enumerate(runs):
        start = len(full_text)
        full_text += run.text or ""
        end = len(full_text)
        run_boundaries.append((start, end, i))

    match_start = full_text.find(find)
    if match_start == -1:
        return False

    match_end = match_start + len(find)

    # Find which runs are involved
    first_run_idx = None
    last_run_idx = None
    for start, end, idx in run_boundaries:
        if start < match_end and end > match_start:
            if first_run_idx is None:
                first_run_idx = idx
            last_run_idx = idx

    if first_run_idx is None:
        return False

    # Put the replacement text in the first run, trimming the
    # matched portion. Keep text before the match in the first run
    # and text after the match in the last run.
    first_run = runs[first_run_idx]
    last_run = runs[last_run_idx]

    first_run_start = run_boundaries[first_run_idx][0]
    last_run_end = run_boundaries[last_run_idx][1]

    prefix = full_text[first_run_start:match_start]
    suffix = full_text[match_end:last_run_end]

    first_run.text = prefix + replace + (
        suffix if first_run_idx == last_run_idx else ""
    )

    # Clear intermediate and last runs that were part of the match
    for idx in range(first_run_idx + 1, last_run_idx + 1):
        if idx == last_run_idx and first_run_idx != last_run_idx:
            # Last run keeps its suffix (text after the match)
            run_start = run_boundaries[idx][0]
            runs[idx].text = full_text[match_end:run_boundaries[idx][1]]
        else:
            runs[idx].text = ""

    return True


def _apply_substitutions_to_docx(
    docx_bytes: bytes,
    substitutions: list[dict[str, str]],
) -> tuple[bytes, list[str], list[str]]:
    """Apply find-and-replace substitutions to a DOCX file.

    Returns (docx_bytes, matched_finds, unmatched_finds).
    """
    from docx import Document

    doc = Document(BytesIO(docx_bytes))
    matched: set[str] = set()

    def process_paragraphs(paragraphs: Any) -> None:
        for paragraph in paragraphs:
            for sub in substitutions:
                if _replace_in_paragraph(
                    paragraph, sub["find"], sub["replace"]
                ):
                    matched.add(sub["find"])

    process_paragraphs(doc.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)

    for section in doc.sections:
        for header in (section.header, section.first_page_header):
            if header and header.is_linked_to_previous is False:
                process_paragraphs(header.paragraphs)
        for footer in (section.footer, section.first_page_footer):
            if footer and footer.is_linked_to_previous is False:
                process_paragraphs(footer.paragraphs)

    buf = BytesIO()
    doc.save(buf)

    all_finds = [s["find"] for s in substitutions]
    unmatched = [f for f in all_finds if f not in matched]
    return buf.getvalue(), sorted(matched), unmatched


def _add_table_loop_to_docx(
    docx_bytes: bytes,
    table_index: int,
    header_rows: int,
    loop_var: str,
    collection: str,
    columns: list[str],
) -> bytes:
    """Convert a table's data rows into a Jinja2 for-loop.

    Keeps header rows, removes all data rows, and inserts a single
    template row wrapped in {%tr for %} / {%tr endfor %} tags.
    Uses docxtpl's {%tr %} syntax for table-row loops.
    """
    from docx import Document
    from copy import deepcopy
    from lxml import etree

    doc = Document(BytesIO(docx_bytes))

    if table_index >= len(doc.tables):
        raise ValueError(
            f"Table index {table_index} out of range "
            f"(document has {len(doc.tables)} tables)"
        )

    table = doc.tables[table_index]
    tbl_elem = table._tbl
    rows = list(tbl_elem)

    # Separate header rows from data rows
    # Filter to only <w:tr> elements
    nsmap = tbl_elem.nsmap
    w_ns = nsmap.get("w", "http://schemas.openxmlformats.org/wordprocessingml/2006/main")
    tr_tag = f"{{{w_ns}}}tr"
    tc_tag = f"{{{w_ns}}}tc"
    p_tag = f"{{{w_ns}}}p"
    r_tag = f"{{{w_ns}}}r"
    t_tag = f"{{{w_ns}}}t"

    tr_elements = [r for r in rows if r.tag == tr_tag]

    if header_rows >= len(tr_elements):
        raise ValueError(
            f"header_rows={header_rows} but table only has "
            f"{len(tr_elements)} rows"
        )

    if len(columns) == 0:
        raise ValueError("columns list cannot be empty")

    # Use the first data row as template for formatting
    first_data_row = tr_elements[header_rows]

    # Remove all data rows (keep headers)
    for tr in tr_elements[header_rows:]:
        tbl_elem.remove(tr)

    # Helper to create a row with text in cells
    def make_row(cell_texts: list[str]) -> Any:
        new_tr = deepcopy(first_data_row)
        cells = list(new_tr.iter(tc_tag))
        for i, text in enumerate(cell_texts):
            if i >= len(cells):
                break
            # Clear existing content in cell
            for p in list(cells[i].iter(p_tag)):
                cells[i].remove(p)
            # Add new paragraph with text
            new_p = etree.SubElement(cells[i], p_tag)
            new_r = etree.SubElement(new_p, r_tag)
            new_t = etree.SubElement(new_r, t_tag)
            new_t.set(f"{{{w_ns}}}space", "preserve")
            new_t.text = text
        return new_tr

    # Build 3 rows using docxtpl's {%tr %} row-loop syntax:
    # Row 1: {%tr for ... %} in its own row (removed during render)
    # Row 2: data template row (repeated for each item)
    # Row 3: {%tr endfor %} in its own row (removed during render)
    for_tag = f"{{%tr for {loop_var} in {collection} %}}"
    endfor_tag = "{%tr endfor %}"

    open_cells = [for_tag] + [""] * (len(columns) - 1)
    open_row = make_row(open_cells)

    data_row = make_row(columns)

    close_cells = [endfor_tag] + [""] * (len(columns) - 1)
    close_row = make_row(close_cells)

    tbl_elem.append(open_row)
    tbl_elem.append(data_row)
    tbl_elem.append(close_row)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _modify_table_in_docx(
    docx_bytes: bytes,
    table_index: int,
    remove_columns: list[int] | None = None,
    rename_headers: dict[int, str] | None = None,
    remove_entire: bool = False,
) -> bytes:
    """Modify a table's structure.

    Args:
        docx_bytes: Source DOCX bytes.
        table_index: 0-based index of the table.
        remove_columns: List of 0-based column indices to remove.
        rename_headers: Dict of {column_index: new_header_text}.
        remove_entire: If True, remove the entire table from the doc.
    """
    from docx import Document

    doc = Document(BytesIO(docx_bytes))

    if table_index >= len(doc.tables):
        raise ValueError(
            f"Table index {table_index} out of range "
            f"(document has {len(doc.tables)} tables)"
        )

    table = doc.tables[table_index]

    if remove_entire:
        table._tbl.getparent().remove(table._tbl)
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # Rename headers first (before removing columns shifts indices)
    if rename_headers:
        header_row = table.rows[0]
        for col_idx, new_text in rename_headers.items():
            if col_idx < len(header_row.cells):
                cell = header_row.cells[col_idx]
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = ""
                    if paragraph.runs:
                        paragraph.runs[0].text = new_text
                    elif paragraph.text == "":
                        paragraph.text = new_text

    # Remove columns (process in reverse order to preserve indices)
    if remove_columns:
        for col_idx in sorted(remove_columns, reverse=True):
            for row in table.rows:
                cells = row.cells
                if col_idx < len(cells):
                    tc = cells[col_idx]._tc
                    tc.getparent().remove(tc)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _remove_section_from_docx(
    docx_bytes: bytes,
    heading_text: str,
) -> bytes:
    """Remove a section (heading + all content until the next heading).

    Also renumbers remaining sections if they use "N. Title" format.
    """
    from docx import Document

    doc = Document(BytesIO(docx_bytes))
    body = doc.element.body

    # Find the heading paragraph and collect elements to remove
    removing = False
    heading_level = None
    elements_to_remove = []

    for element in list(body):
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            # Check if this is a heading
            style = None
            for ppr in element.iter():
                ppr_tag = ppr.tag.split("}")[-1] if "}" in ppr.tag else ppr.tag
                if ppr_tag == "pStyle":
                    style = ppr.get(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val",
                        ""
                    )

            # Get paragraph text
            para_text = ""
            for t_elem in element.iter():
                t_tag = t_elem.tag.split("}")[-1] if "}" in t_elem.tag else t_elem.tag
                if t_tag == "t" and t_elem.text:
                    para_text += t_elem.text

            if style and style.startswith("Heading"):
                if removing:
                    # Hit the next heading — stop removing
                    removing = False
                if heading_text.lower() in para_text.lower():
                    removing = True
                    heading_level = style

        if removing:
            elements_to_remove.append(element)

    for elem in elements_to_remove:
        body.remove(elem)

    # Renumber sections with "N. Title" pattern
    section_num = 1
    for element in body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
        if tag != "p":
            continue
        style = None
        for ppr in element.iter():
            ppr_tag = ppr.tag.split("}")[-1] if "}" in ppr.tag else ppr.tag
            if ppr_tag == "pStyle":
                style = ppr.get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val",
                    ""
                )
        if not style or not style.startswith("Heading"):
            continue

        # Find the text run and renumber if it matches "N. Title"
        for t_elem in element.iter():
            t_tag = t_elem.tag.split("}")[-1] if "}" in t_elem.tag else t_elem.tag
            if t_tag == "t" and t_elem.text:
                import re as _re
                match = _re.match(r"^\d+\.\s+", t_elem.text)
                if match:
                    t_elem.text = _re.sub(
                        r"^\d+\.\s+",
                        f"{section_num}. ",
                        t_elem.text,
                    )
                    section_num += 1
                    break

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_content_after_heading(
    docx_bytes: bytes,
    heading_text: str,
    paragraphs: list[dict[str, str]],
) -> bytes:
    """Insert paragraphs after a heading in the document.

    Each paragraph dict has "text" and optional "style" (defaults to
    "Normal"). Removes any existing content between the heading and
    the next heading before inserting.

    Args:
        docx_bytes: Source DOCX.
        heading_text: Heading to insert after (case-insensitive match).
        paragraphs: List of {"text": "...", "style": "Normal"} dicts.
    """
    from docx import Document
    from lxml import etree

    doc = Document(BytesIO(docx_bytes))
    body = doc.element.body

    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    p_tag = f"{{{w_ns}}}p"
    r_tag = f"{{{w_ns}}}r"
    t_tag = f"{{{w_ns}}}t"
    ppr_tag = f"{{{w_ns}}}pPr"
    pstyle_tag = f"{{{w_ns}}}pStyle"

    # Find the heading element
    heading_elem = None
    for element in body:
        if element.tag != p_tag:
            continue
        # Get style
        style = None
        ppr = element.find(ppr_tag)
        if ppr is not None:
            ps = ppr.find(pstyle_tag)
            if ps is not None:
                style = ps.get(f"{{{w_ns}}}val", "")
        if not style or not style.startswith("Heading"):
            continue
        # Get text
        para_text = ""
        for t_elem in element.iter(t_tag):
            if t_elem.text:
                para_text += t_elem.text
        if heading_text.lower() in para_text.lower():
            heading_elem = element
            break

    if heading_elem is None:
        raise ValueError(f"Heading '{heading_text}' not found")

    # Remove existing content between this heading and the next
    removing = False
    to_remove = []
    for element in list(body):
        if element is heading_elem:
            removing = True
            continue
        if not removing:
            continue
        # Stop at next heading or table
        if element.tag == p_tag:
            ppr = element.find(ppr_tag)
            if ppr is not None:
                ps = ppr.find(pstyle_tag)
                if ps is not None:
                    val = ps.get(f"{{{w_ns}}}val", "")
                    if val.startswith("Heading"):
                        break
        to_remove.append(element)

    for elem in to_remove:
        body.remove(elem)

    # Find insertion point (right after heading)
    heading_index = list(body).index(heading_elem)

    # Build and insert new paragraphs
    # Get a style mapping from existing doc styles
    for i, para_def in enumerate(paragraphs):
        text = para_def.get("text", "")
        style = para_def.get("style", "Normal")

        new_p = etree.SubElement(body, p_tag)
        # Move to correct position
        body.remove(new_p)
        body.insert(heading_index + 1 + i, new_p)

        # Add style
        new_ppr = etree.SubElement(new_p, ppr_tag)
        new_pstyle = etree.SubElement(new_ppr, pstyle_tag)
        new_pstyle.set(f"{{{w_ns}}}val", style)

        # Add text run
        new_r = etree.SubElement(new_p, r_tag)
        new_t = etree.SubElement(new_r, t_tag)
        new_t.set(f"{{{w_ns}}}space", "preserve")
        new_t.text = text

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Tool Functions ──


def _check_tool_limit(deps: ConversionDeps, tool_name: str) -> str | None:
    """Increment counter and return error string if over limit."""
    deps.tool_call_count += 1
    if deps.tool_call_count > deps.max_tool_calls:
        msg = TOOL_LIMIT_MSG.format(limit=deps.max_tool_calls)
        deps.event_stream.push("error", {"message": msg})
        return msg
    return None


def _trace_tool(
    deps: ConversionDeps,
    tool: str,
    seq: int,
    status: str,
    tool_input: Any,
    output: str,
    duration_ms: int,
) -> None:
    """Append a tool call entry to the trace log."""
    deps.trace_entries.append({
        "sequence": seq,
        "tool": tool,
        "status": status,
        "input": tool_input,
        "output": output[:1000],
        "duration_ms": duration_ms,
        "timestamp": time.time(),
    })


async def apply_substitutions_tool(
    ctx: RunContext[ConversionDeps],
    substitutions: list[dict[str, str]],
) -> str:
    """Apply find-and-replace substitutions to the original DOCX.

    Each substitution is a dict with "find" (the literal text in the
    original document) and "replace" (the Jinja2 placeholder to use,
    e.g. "{{ operator_name }}").

    This preserves all original formatting — fonts, tables, styles,
    centering, etc. Only the text content changes.
    """
    deps = ctx.deps
    seq = deps.tool_call_count + 1

    err = _check_tool_limit(deps, "apply_substitutions")
    if err:
        return err

    t0 = time.time()
    tool_input = {"substitutions": substitutions}

    deps.event_stream.push(
        "tool_call",
        {
            "tool": "apply_substitutions",
            "status": "running",
            "sequence": seq,
        },
    )

    try:
        # Read from template.docx if it exists (preserves prior
        # modifications like table loops), otherwise from original
        source = (
            "template.docx"
            if deps.state.exists("template.docx")
            else "original"
        )
        source_bytes = deps.state.read(source)
        result_bytes, matched, unmatched = _apply_substitutions_to_docx(
            source_bytes, substitutions
        )
        deps.state.write("template.docx", result_bytes)

        n_total = len(substitutions)
        n_matched = len(matched)
        n_unmatched = len(unmatched)

        if n_matched == 0:
            status = "error"
            summary = (
                f"0 of {n_total} substitutions matched. The 'find' "
                f"values must be EXACT substrings from the document. "
                f"Failed: {unmatched[:3]}"
            )
        elif n_unmatched > 0:
            status = "success"
            summary = (
                f"{n_matched} of {n_total} matched. "
                f"Unmatched: {unmatched[:5]}"
            )
        else:
            status = "success"
            summary = f"All {n_matched} substitutions matched."

        deps.event_stream.push(
            "tool_result",
            {
                "tool": "apply_substitutions",
                "status": status,
                "sequence": seq,
                "summary": summary[:200],
            },
        )

        result_msg = f"{n_matched} of {n_total} substitutions matched."
        if unmatched:
            result_msg += (
                f"\n\nWARNING: {n_unmatched} substitutions did NOT "
                f"match any text in the document. These 'find' values "
                f"were not found (check for exact spelling, spacing, "
                f"and punctuation):\n"
                + "\n".join(f'  - "{u}"' for u in unmatched[:10])
            )
        result_msg += "\nfile_id=template.docx"
        _trace_tool(deps, "apply_substitutions", seq, status,
                     tool_input, result_msg,
                     int((time.time() - t0) * 1000))
        return result_msg
    except Exception as e:
        err_msg = f"Error applying substitutions: {e}"
        _trace_tool(deps, "apply_substitutions", seq, "error",
                     tool_input, err_msg,
                     int((time.time() - t0) * 1000))
        deps.event_stream.push(
            "tool_result",
            {
                "tool": "apply_substitutions",
                "status": "error",
                "sequence": seq,
                "summary": str(e)[:200],
            },
        )
        return err_msg


async def add_table_loop_tool(
    ctx: RunContext[ConversionDeps],
    table_index: int,
    header_rows: int,
    loop_var: str,
    collection: str,
    columns: list[str],
) -> str:
    """Convert a table's repeating data rows into a Jinja2 for-loop.

    Use this for tables that have repeating rows (e.g., procedure steps,
    materials lists, test results). The header rows are kept as-is, and
    the data rows are replaced with a single template row inside a
    for-loop.

    Args:
        table_index: 0-based index of the table in the document.
        header_rows: Number of header rows to keep (usually 1).
        loop_var: Loop variable name (e.g., "step").
        collection: Collection to iterate (e.g., "steps").
        columns: Jinja2 expressions for each column in the data row
            (e.g., ["{{ step.number }}", "{{ step.description }}"]).
    """
    deps = ctx.deps
    seq = deps.tool_call_count + 1
    t0 = time.time()
    tool_input = {
        "table_index": table_index, "header_rows": header_rows,
        "loop_var": loop_var, "collection": collection,
        "columns": columns,
    }

    err = _check_tool_limit(deps, "add_table_loop")
    if err:
        return err

    deps.event_stream.push(
        "tool_call",
        {
            "tool": "add_table_loop",
            "status": "running",
            "sequence": seq,
        },
    )

    try:
        source = (
            "template.docx"
            if deps.state.exists("template.docx")
            else "original"
        )
        source_bytes = deps.state.read(source)
        result_bytes = _add_table_loop_to_docx(
            source_bytes,
            table_index,
            header_rows,
            loop_var,
            collection,
            columns,
        )
        deps.state.write("template.docx", result_bytes)

        summary = (
            f"Table {table_index}: {{% for {loop_var} in "
            f"{collection} %}} with {len(columns)} columns"
        )
        deps.event_stream.push(
            "tool_result",
            {
                "tool": "add_table_loop",
                "status": "success",
                "sequence": seq,
                "summary": summary[:200],
            },
        )
        result_msg = (
            f"Success. Table {table_index} now has a for-loop: "
            f"{{% for {loop_var} in {collection} %}}. "
            f"file_id=template.docx"
        )
        _trace_tool(deps, "add_table_loop", seq, "success",
                     tool_input, result_msg,
                     int((time.time() - t0) * 1000))
        return result_msg
    except Exception as e:
        err_msg = f"Error adding table loop: {e}"
        _trace_tool(deps, "add_table_loop", seq, "error",
                     tool_input, err_msg,
                     int((time.time() - t0) * 1000))
        deps.event_stream.push(
            "tool_result",
            {
                "tool": "add_table_loop",
                "status": "error",
                "sequence": seq,
                "summary": str(e)[:200],
            },
        )
        return err_msg


async def modify_table_tool(
    ctx: RunContext[ConversionDeps],
    table_index: int,
    remove_columns: list[int] | None = None,
    rename_headers: dict[int, str] | None = None,
    remove_entire: bool = False,
) -> str:
    """Modify a table — remove columns, rename headers, or delete it.

    Args:
        table_index: 0-based index of the table in the document.
        remove_columns: List of 0-based column indices to remove.
        rename_headers: Dict mapping column index to new header text.
        remove_entire: If True, remove the entire table.
    """
    deps = ctx.deps
    seq = deps.tool_call_count + 1
    t0 = time.time()
    tool_input = {
        "table_index": table_index,
        "remove_columns": remove_columns,
        "rename_headers": rename_headers,
        "remove_entire": remove_entire,
    }

    err = _check_tool_limit(deps, "modify_table")
    if err:
        return err

    deps.event_stream.push(
        "tool_call",
        {
            "tool": "modify_table",
            "status": "running",
            "sequence": seq,
        },
    )

    try:
        source = (
            "template.docx"
            if deps.state.exists("template.docx")
            else "original"
        )
        source_bytes = deps.state.read(source)
        result_bytes = _modify_table_in_docx(
            source_bytes,
            table_index,
            remove_columns=remove_columns,
            rename_headers=rename_headers,
            remove_entire=remove_entire,
        )
        deps.state.write("template.docx", result_bytes)

        parts = []
        if remove_entire:
            parts.append("removed entire table")
        if remove_columns:
            parts.append(f"removed columns {remove_columns}")
        if rename_headers:
            parts.append(
                f"renamed {len(rename_headers)} header(s)"
            )
        summary = f"Table {table_index}: {', '.join(parts)}"

        deps.event_stream.push(
            "tool_result",
            {
                "tool": "modify_table",
                "status": "success",
                "sequence": seq,
                "summary": summary[:200],
            },
        )
        result_msg = f"Success. {summary}. file_id=template.docx"
        _trace_tool(deps, "modify_table", seq, "success",
                     tool_input, result_msg,
                     int((time.time() - t0) * 1000))
        return result_msg
    except Exception as e:
        err_msg = f"Error modifying table: {e}"
        _trace_tool(deps, "modify_table", seq, "error",
                     tool_input, err_msg,
                     int((time.time() - t0) * 1000))
        deps.event_stream.push(
            "tool_result",
            {
                "tool": "modify_table",
                "status": "error",
                "sequence": seq,
                "summary": str(e)[:200],
            },
        )
        return err_msg


async def remove_section_tool(
    ctx: RunContext[ConversionDeps],
    heading_text: str,
) -> str:
    """Remove a section by its heading text.

    Removes the heading paragraph and all content until the next
    heading of the same or higher level. Also renumbers remaining
    sections if they use "N. Title" format (e.g., "1. Purpose",
    "2. Scope").

    Args:
        heading_text: The heading text to match (case-insensitive
            partial match). E.g., "Results" matches "5. Results".
    """
    deps = ctx.deps
    seq = deps.tool_call_count + 1
    t0 = time.time()
    tool_input = {"heading_text": heading_text}

    err = _check_tool_limit(deps, "remove_section")
    if err:
        return err

    deps.event_stream.push(
        "tool_call",
        {
            "tool": "remove_section",
            "status": "running",
            "sequence": seq,
        },
    )

    try:
        source = (
            "template.docx"
            if deps.state.exists("template.docx")
            else "original"
        )
        source_bytes = deps.state.read(source)
        result_bytes = _remove_section_from_docx(
            source_bytes, heading_text
        )
        deps.state.write("template.docx", result_bytes)

        summary = (
            f"Removed section '{heading_text}' and renumbered"
        )
        deps.event_stream.push(
            "tool_result",
            {
                "tool": "remove_section",
                "status": "success",
                "sequence": seq,
                "summary": summary[:200],
            },
        )
        result_msg = f"Success. {summary}. file_id=template.docx"
        _trace_tool(deps, "remove_section", seq, "success",
                     tool_input, result_msg,
                     int((time.time() - t0) * 1000))
        return result_msg
    except Exception as e:
        err_msg = f"Error removing section: {e}"
        _trace_tool(deps, "remove_section", seq, "error",
                     tool_input, err_msg,
                     int((time.time() - t0) * 1000))
        deps.event_stream.push(
            "tool_result",
            {
                "tool": "remove_section",
                "status": "error",
                "sequence": seq,
                "summary": str(e)[:200],
            },
        )
        return err_msg


async def add_content_tool(
    ctx: RunContext[ConversionDeps],
    after_heading: str,
    paragraphs: list[dict[str, str]],
) -> str:
    """Insert content after a section heading.

    Replaces any existing content between the heading and the next
    heading with the provided paragraphs. Use this to add structured
    content like numbered lists, bullet points, or formatted text.

    Args:
        after_heading: Heading text to insert after (partial match).
        paragraphs: List of dicts with "text" and optional "style".
            Style can be "Normal", "List Number", "List Bullet", etc.
            Jinja2 syntax in text is preserved as-is.

    Example — converting a section to a numbered list:
        add_content(
            after_heading="Procedure",
            paragraphs=[
                {"text": "{% for step in steps %}", "style": "Normal"},
                {"text": "{{ step.name }}: {{ step.description }}",
                 "style": "List Number"},
                {"text": "{% endfor %}", "style": "Normal"},
            ]
        )
    """
    deps = ctx.deps
    seq = deps.tool_call_count + 1
    t0 = time.time()
    tool_input = {
        "after_heading": after_heading,
        "paragraphs": paragraphs,
    }

    err = _check_tool_limit(deps, "add_content")
    if err:
        return err

    deps.event_stream.push(
        "tool_call",
        {
            "tool": "add_content",
            "status": "running",
            "sequence": seq,
        },
    )

    try:
        source = (
            "template.docx"
            if deps.state.exists("template.docx")
            else "original"
        )
        source_bytes = deps.state.read(source)
        result_bytes = _add_content_after_heading(
            source_bytes, after_heading, paragraphs
        )
        deps.state.write("template.docx", result_bytes)

        summary = (
            f"Added {len(paragraphs)} paragraph(s) after "
            f"'{after_heading}'"
        )
        deps.event_stream.push(
            "tool_result",
            {
                "tool": "add_content",
                "status": "success",
                "sequence": seq,
                "summary": summary[:200],
            },
        )
        result_msg = f"Success. {summary}. file_id=template.docx"
        _trace_tool(deps, "add_content", seq, "success",
                     tool_input, result_msg,
                     int((time.time() - t0) * 1000))
        return result_msg
    except Exception as e:
        err_msg = f"Error adding content: {e}"
        _trace_tool(deps, "add_content", seq, "error",
                     tool_input, err_msg,
                     int((time.time() - t0) * 1000))
        deps.event_stream.push(
            "tool_result",
            {
                "tool": "add_content",
                "status": "error",
                "sequence": seq,
                "summary": str(e)[:200],
            },
        )
        return err_msg


async def validate_tool(
    ctx: RunContext[ConversionDeps], file_id: str
) -> str:
    """Validate the template DOCX and check Jinja2 syntax.

    Renders the template with mock data, checks for render errors and
    surviving Jinja2 syntax, and extracts detected variables.
    """
    deps = ctx.deps
    seq = deps.tool_call_count + 1
    t0 = time.time()
    tool_input = {"file_id": file_id}

    err = _check_tool_limit(deps, "validate")
    if err:
        return err

    deps.event_stream.push(
        "tool_call",
        {"tool": "validate", "status": "running", "sequence": seq},
    )

    try:
        template_bytes = deps.state.read(file_id)
        render_result = _try_render(template_bytes, deps.template_type)

        # Always save preview PDF if we got one
        if render_result.pdf_bytes:
            deps.state.write("preview.pdf", render_result.pdf_bytes)

        all_text = _extract_docx_text(template_bytes)
        detected_vars = _extract_jinja_variables(all_text)

        # Known loop variables and their collections
        known_loop_vars = {
            "step", "role", "note", "figure",
            "non_image_attachment",
        }
        known_collections = {
            "steps", "roles", "notes", "figures",
            "non_image_attachments",
        }
        unknown_vars = sorted(
            v for v in detected_vars
            if v not in KNOWN_VARIABLES
            and v not in known_loop_vars
            and v not in known_collections
        )

        # Check for invalid dotted attributes (e.g., step.operator)
        known_attrs = {
            "step": {"name", "description", "duration_min", "role_name"},
            "role": {"name"},
            "note": {"content", "author_name", "created_at"},
            "figure": {"filename", "number"},
        }
        dotted_pattern = re.compile(r"\{\{\s*([\w]+)\.([\w]+)\s*\}\}")
        invalid_attrs: list[str] = []
        for match in dotted_pattern.finditer(all_text):
            obj, attr = match.group(1), match.group(2)
            if obj in known_attrs and attr not in known_attrs[obj]:
                invalid_attrs.append(f"{obj}.{attr}")
        invalid_attrs = sorted(set(invalid_attrs))

        # QA: check that a steps loop exists
        has_steps_loop = "{% for step in steps %}" in all_text or \
            "{%tr for step in steps %}" in all_text

        issues = []
        if render_result.render_error:
            issues.append(f"Render error: {render_result.render_error}")
        if render_result.jinja_remnants:
            issues.append(
                f"Jinja remnants after render: "
                f"{render_result.jinja_remnants}"
            )
        if unknown_vars:
            issues.append(
                f"Unknown variables (not supported by the system): "
                f"{unknown_vars}. These are NOT valid variables. "
                f"Call apply_substitutions to replace each one back "
                f"to its original text from the document. For "
                f"example, if you substituted 'Dr. Sarah Chen' with "
                f"'{{{{ prepared_by }}}}', call apply_substitutions "
                f"with find='{{{{ prepared_by }}}}' and "
                f"replace='Dr. Sarah Chen' to revert it."
            )
        if invalid_attrs:
            issues.append(
                f"Invalid attributes (don't exist in the system): "
                f"{invalid_attrs}. Remove these from the template "
                f"using modify_table to remove the columns, or "
                f"replace with valid attributes."
            )
        if not has_steps_loop:
            issues.append(
                "QA: No {{% for step in steps %}} loop was found. "
                "Every SOP/batch record template MUST have a steps "
                "loop for the procedure section. Look at the "
                "original document — find the table with repeating "
                "process/procedure rows and convert it using "
                "add_table_loop with collection='steps'."
            )

        status_str = "error" if issues else "success"
        summary = (
            "; ".join(issues)
            if issues
            else f"Valid. {len(detected_vars)} variables found."
        )

        deps.event_stream.push(
            "tool_result",
            {
                "tool": "validate",
                "status": status_str,
                "sequence": seq,
                "summary": summary[:200],
            },
        )

        result_parts = [f"Variables found: {sorted(detected_vars)}"]
        if issues:
            result_parts.append(f"Issues: {'; '.join(issues)}")
        else:
            result_parts.append(
                "No issues found. Template renders cleanly."
            )
        result_msg = "\n".join(result_parts)
        _trace_tool(deps, "validate", seq, status_str,
                     tool_input, result_msg,
                     int((time.time() - t0) * 1000))
        return result_msg
    except Exception as e:
        err_msg = f"Error validating: {e}"
        _trace_tool(deps, "validate", seq, "error",
                     tool_input, err_msg,
                     int((time.time() - t0) * 1000))
        deps.event_stream.push(
            "tool_result",
            {
                "tool": "validate",
                "status": "error",
                "sequence": seq,
                "summary": str(e)[:200],
            },
        )
        return err_msg


async def compare_to_original_tool(
    ctx: RunContext[ConversionDeps], file_id: str
) -> str:
    """Render the template with mock data and compare to the original.

    Renders the template to PDF, then spawns a vision subagent to
    compare the rendered output against the original document.
    Returns a text assessment of visual fidelity.
    """
    deps = ctx.deps
    seq = deps.tool_call_count + 1
    t0 = time.time()
    tool_input = {"file_id": file_id}

    err = _check_tool_limit(deps, "compare_to_original")
    if err:
        return err

    deps.event_stream.push(
        "tool_call",
        {
            "tool": "compare_to_original",
            "status": "running",
            "sequence": seq,
        },
    )

    try:
        template_bytes = deps.state.read(file_id)
        render_result = _try_render(template_bytes, deps.template_type)

        if not render_result.pdf_bytes:
            deps.event_stream.push(
                "tool_result",
                {
                    "tool": "compare_to_original",
                    "status": "error",
                    "sequence": seq,
                    "summary": (
                        "Could not render PDF for comparison"
                        f" ({render_result.render_error})"
                    ),
                },
            )
            return (
                f"Cannot compare: render failed "
                f"({render_result.render_error})"
            )

        deps.state.write("preview.pdf", render_result.pdf_bytes)

        comparison_prompt = (
            "Compare these two documents. The first is the ORIGINAL "
            "filled document. The second is a RENDERED TEMPLATE with "
            "mock data filled in.\n\n"
            "Assess:\n"
            "1. Structural match — sections, tables, headings in same "
            "order?\n"
            "2. Formatting similarity — fonts, spacing, layout\n"
            "3. Missing or extra content\n\n"
            "Be specific about differences. If they match well, say so "
            "briefly."
        )

        original_bytes = deps.original_pdf_bytes
        original_mime = "application/pdf"
        if original_bytes[:4] == b"\x89PNG":
            original_mime = "image/png"
        elif original_bytes[:2] == b"\xff\xd8":
            original_mime = "image/jpeg"

        vision_agent: Agent[None, str] = Agent(
            deps.model,
            system_prompt="You are a document comparison assistant.",
            output_type=str,
        )

        user_content: list[Any] = [
            BinaryContent(
                data=original_bytes, media_type=original_mime
            ),
            BinaryContent(
                data=render_result.pdf_bytes,
                media_type="application/pdf",
            ),
            comparison_prompt,
        ]

        vision_result = await asyncio.wait_for(
            vision_agent.run(user_content),
            timeout=120,
        )

        assessment = vision_result.output
        deps.event_stream.push(
            "tool_result",
            {
                "tool": "compare_to_original",
                "status": "success",
                "sequence": seq,
                "summary": assessment[:200],
            },
        )
        _trace_tool(deps, "compare_to_original", seq, "success",
                     tool_input, assessment,
                     int((time.time() - t0) * 1000))
        return assessment
    except Exception as e:
        err_msg = f"Comparison failed: {e}"
        _trace_tool(deps, "compare_to_original", seq, "error",
                     tool_input, err_msg,
                     int((time.time() - t0) * 1000))
        deps.event_stream.push(
            "tool_result",
            {
                "tool": "compare_to_original",
                "status": "error",
                "sequence": seq,
                "summary": str(e)[:200],
            },
        )
        return err_msg


# ── System Prompt ──


SYSTEM_PROMPT = """\
You are a document template engineer. You convert filled SOPs and batch \
records into reusable Jinja2 templates.

You have 7 tools:
1. apply_substitutions(substitutions) — Replace filled values with \
Jinja2 placeholders
2. add_table_loop(table_index, header_rows, loop_var, collection, \
columns) — Convert a table's repeating rows into a for-loop
3. modify_table(table_index, remove_columns, rename_headers, \
remove_entire) — Remove columns, rename headers, or delete a table
4. remove_section(heading_text) — Remove a section by heading \
(heading + content + auto-renumber)
5. add_content(after_heading, paragraphs) — Insert properly styled \
paragraphs after a heading. Use for lists, formatted text, etc.
6. validate(file_id) — Render the template with mock data and check \
for errors
7. compare_to_original(file_id) — Visually compare rendered output

YOUR WORKFLOW:
1. Read the document text carefully
2. Identify single-value fields (names, dates, IDs, etc.) AND \
repeating tables (procedure steps, materials, test results)
3. Call apply_substitutions() for all single-value replacements
4. Call add_table_loop() for EACH table with repeating data rows — \
procedure steps, materials lists, test results, approval tables, etc.
5. Call validate("template.docx") to check for errors
6. If errors: fix and retry
7. Once validation passes, call compare_to_original("template.docx")

SUBSTITUTION FORMAT:
Each substitution has "find" (exact text from the document) and \
"replace" (the Jinja2 placeholder).

Example — ONLY using known variables from the list above:
[
  {{"find": "Acme Therapeutics", "replace": "{{{{ organization_name }}}}"}},
  {{"find": "AAV Campaign Q1", "replace": "{{{{ project_name }}}}"}},
  {{"find": "This SOP describes the preparation of...", \
"replace": "{{{{ protocol_description }}}}"}}
]

IMPORTANT: Do NOT substitute values that have no matching known \
variable. "SOP-BP-2026-042" (a document number), "Dr. Sarah Chen" \
(a person name), "January 15, 2026" (a specific date), \
"Upstream Process Development" (a department) — these do NOT map to \
any known variable and must be LEFT AS-IS in the document.

TABLE LOOPS — ONLY for these collections:
- "steps" (loop_var="step"): procedure/process step tables
- "roles" (loop_var="role"): role assignment or sign-off tables

Do NOT loop any other tables. Materials, equipment, results, \
metadata tables are STATIC content — leave them as-is with their \
original values. They are specific to each document.

Example — a procedure steps table (table index 2, 1 header row):
add_table_loop(
  table_index=2,
  header_rows=1,
  loop_var="step",
  collection="steps",
  columns=["{{{{ step.name }}}}", "{{{{ step.description }}}}", \
"{{{{ step.duration_min }}}}"]
)

IMPORTANT:
- Every SOP/batch record MUST have a steps loop. Find the table in \
the document that lists the procedure steps (it may be labeled \
"Procedure", "Process Steps", "Operations", or similar) and \
convert it with add_table_loop using collection="steps".
- Call add_table_loop AFTER apply_substitutions. Tables you convert \
to loops should NOT have their cell values substituted.

VARIABLE DEFINITIONS — only substitute values that clearly match:

VALUE-LEVEL variables (replace a specific value within a line):
- protocol_name: the SOP/protocol TITLE text (e.g., "Buffer Preparation SOP")
- version_number: a version number (e.g., "3", "v2.1")
- created_at: document creation date
- organization_name: the COMPANY/ORG name (e.g., "Acme Therapeutics")
- project_name: the project/campaign name (e.g., "AAV Campaign Q1")
- run_name: a run identifier (e.g., "Run-2026-001")
- run_status: ONLY status words like "COMPLETED", "ACTIVE", "DRAFT"
- started_at: run start timestamp
- completed_at: run end timestamp

PARAGRAPH-LEVEL variables (replace an ENTIRE paragraph or sentence):
- protocol_description: replaces the ENTIRE purpose/description \
paragraph. Find the FULL sentence that describes what this SOP is \
for and replace ALL of it. For example, find the complete text \
"This SOP describes the preparation of Tris-HCl buffer (50 mM, \
pH 7.4) for use in downstream purification steps of monoclonal \
antibody production at Acme Therapeutics." and replace the whole \
thing with {{{{ protocol_description }}}}.

For table loops:
  step: step.name, step.description, step.duration_min, step.role_name
  role: role.name
  note: note.content, note.author_name, note.created_at
  figure: figure.filename, figure.number

CRITICAL RULES:
- The "find" value must be an EXACT substring from the document.
- WHEN IN DOUBT, LEAVE THE VALUE AS-IS. Most values in a filled \
document are static content (measurements, person names, lot \
numbers, equipment IDs, results). Only substitute when you are \
CERTAIN the value maps to a known variable above.
- Person names (e.g., "Dr. Sarah Chen") are NOT organization_name. \
Leave them as-is.
- Measurements (pH, volumes, conductivity) are NOT run_status or \
version_number. Leave them as-is.
- Department names are NOT protocol_description. Leave them as-is.
- Document numbers (e.g., "SOP-BP-2026-042") are NOT protocol_name. \
Leave them as-is.
- Lot numbers, equipment IDs, specifications — leave as-is.
- Do NOT substitute values in tables you plan to convert to loops.
- Do NOT substitute static text (section headers, instructions, labels).
- Keep formatting labels like "Document Number:", "Prepared By:" intact.

You have a limited number of tool calls. If a tool tells you the limit \
is reached, stop immediately and report what you have so far.
""".format(known_vars=", ".join(sorted(KNOWN_VARIABLES)))


def _save_trace(
    deps: ConversionDeps,
    action: str,
    filename: str,
    model_str: str,
    outcome: str,
    error_msg: str | None = None,
) -> None:
    """Write a trace.json log of the entire conversion/refinement."""
    from datetime import datetime, timezone

    trace = {
        "conversion_id": str(deps.state.conversion_id),
        "action": action,
        "started_at": datetime.fromtimestamp(
            deps.started_at, tz=timezone.utc
        ).isoformat(),
        "completed_at": datetime.now(timezone.utc).isoformat(),
        "duration_seconds": round(time.time() - deps.started_at, 2),
        "model": model_str,
        "org_id": str(deps.org_id),
        "template_type": deps.template_type,
        "input_filename": filename,
        "outcome": outcome,
        "tool_calls_used": deps.tool_call_count,
        "tool_calls_limit": deps.max_tool_calls,
        "tool_calls": deps.trace_entries,
        "error": error_msg,
    }

    # Append to trace log (supports multiple actions per conversion)
    trace_file = "trace.json"
    existing: list = []
    if deps.state.exists(trace_file):
        try:
            loaded = deps.state.read_json(trace_file)
            if isinstance(loaded, list):
                existing = loaded
        except Exception:
            pass
    existing.append(trace)
    deps.state.write_json(trace_file, existing)


# ── Main Conversion Pipeline ──


async def convert_document(
    db: AsyncSession,
    org_id: UUID,
    file_bytes: bytes,
    filename: str,
    template_type: str,
    conversion_id: UUID | None = None,
) -> dict[str, Any]:
    """Main conversion: DOCX -> AI identifies variables -> substitutions.

    The AI model analyzes the document text, identifies variable values,
    and calls apply_substitutions to create the template. The original
    DOCX formatting is fully preserved.
    """
    from app.services.ai.ai_config import get_model

    if conversion_id is None:
        conversion_id = str(ULID())

    state = ConversionState(org_id, conversion_id)
    state.ensure_dir()

    event_stream = EventStream()
    _active_streams[str(conversion_id)] = event_stream

    try:
        # Store original
        state.write("original", file_bytes)

        ext = Path(filename).suffix.lower()
        if ext != ".docx":
            raise ValueError(
                "Only DOCX files are supported for template conversion. "
                f"Got: {ext}"
            )

        doc_text = _extract_text_from_docx(file_bytes)
        original_pdf = await _to_pdf(file_bytes, filename)
        state.write("original.pdf", original_pdf)

        # Resolve model
        model = await get_model("template_convert", db, org_id=org_id)

        # Build deps and agent
        deps = ConversionDeps(
            state=state,
            event_stream=event_stream,
            org_id=org_id,
            template_type=template_type,
            original_pdf_bytes=original_pdf,
            model=model,
        )

        agent: Agent[ConversionDeps, str] = Agent(
            model,
            system_prompt=SYSTEM_PROMPT,
            tools=[
                apply_substitutions_tool,
                add_table_loop_tool,
                modify_table_tool,
                remove_section_tool,
                add_content_tool,
                validate_tool,
                compare_to_original_tool,
            ],
            deps_type=ConversionDeps,
        )

        user_prompt = (
            f"Convert this completed {template_type} document into a "
            "Jinja2 template by identifying all filled-in values and "
            "replacing them with placeholders.\n\n"
            "DOCUMENT CONTENT:\n"
            f"```\n{doc_text}\n```\n\n"
            "Identify every value that should become a variable — "
            "names, dates, lot numbers, measurements, IDs, operator "
            "initials, status values, etc. Keep static text (section "
            "headers, instructions, labels) as-is.\n\n"
            "Use apply_substitutions() with all your find/replace "
            "pairs, then validate and compare."
        )

        await asyncio.wait_for(
            agent.run(user_prompt, deps=deps),
            timeout=600,
        )

        # Build final result from state
        detected_vars: set[str] = set()
        if state.exists("template.docx"):
            template_bytes = state.read("template.docx")
            all_text = _extract_docx_text(template_bytes)
            detected_vars = _extract_jinja_variables(all_text)

        # Critical variables that MUST be present
        critical_vars = {"steps"}
        missing_critical = sorted(
            v for v in critical_vars if v not in detected_vars
        )

        warnings: list[dict[str, str]] = []
        for var in missing_critical:
            warnings.append({
                "type": "critical_missing",
                "variable": var,
                "description": (
                    f"CRITICAL: Template is missing '{var}' loop. "
                    "Every SOP/batch record template must have a "
                    "procedure steps loop."
                ),
            })

        missing = sorted(
            v for v in KNOWN_VARIABLES
            if v not in detected_vars and v not in critical_vars
        )
        for var in missing:
            warnings.append({
                "type": "missing_variable",
                "variable": var,
                "description": (
                    f"Template does not use '{var}' — this data "
                    "won't appear in the rendered output"
                ),
            })

        result_data = {
            "conversion_id": str(conversion_id),
            "preview_url": state.preview_url,
            "template_download_url": state.template_url,
            "warnings": warnings,
            "variables_detected": sorted(detected_vars),
        }

        state.write_json("result.json", result_data)

        _save_trace(
            deps, "convert", filename, str(model),
            "success",
        )

        event_stream.push("complete", {
            "template_url": state.template_url,
            "preview_url": (
                state.preview_url
                if state.exists("preview.pdf")
                else None
            ),
            "variables": sorted(detected_vars),
            "warnings": warnings,
        })

        return result_data

    except Exception as e:
        logger.exception("Template conversion failed")
        if deps:
            _save_trace(
                deps, "convert", filename, str(model),
                "error", error_msg=str(e),
            )
        event_stream.push("error", {"message": str(e)})
        raise
    finally:
        event_stream.close()


async def refine_template(
    db: AsyncSession,
    org_id: UUID,
    state: ConversionState,
    instruction: str,
) -> dict[str, Any]:
    """Refine an existing conversion via natural language instruction.

    Uses the same tool-based agent loop as convert_document.
    """
    from app.services.ai.ai_config import get_model

    cid = str(state.conversion_id)
    event_stream = EventStream()
    _active_streams[cid] = event_stream

    try:
        model = await get_model("template_convert", db, org_id=org_id)

        original_pdf = (
            state.read("original.pdf")
            if state.exists("original.pdf")
            else b""
        )

        deps = ConversionDeps(
            state=state,
            event_stream=event_stream,
            org_id=org_id,
            template_type="SOP",
            original_pdf_bytes=original_pdf,
            model=model,
        )

        agent: Agent[ConversionDeps, str] = Agent(
            model,
            system_prompt=SYSTEM_PROMPT,
            tools=[
                apply_substitutions_tool,
                add_table_loop_tool,
                modify_table_tool,
                remove_section_tool,
                add_content_tool,
                validate_tool,
                compare_to_original_tool,
            ],
            deps_type=ConversionDeps,
        )

        # Load current template text and table structure for context
        current_text = ""
        table_summary = ""
        if state.exists("template.docx"):
            template_bytes = state.read("template.docx")
            current_text = _extract_docx_text(template_bytes)

            from docx import Document as _Doc
            doc = _Doc(BytesIO(template_bytes))
            table_lines = []
            for ti, table in enumerate(doc.tables):
                headers = [
                    c.text[:30] for c in table.rows[0].cells
                ]
                table_lines.append(
                    f"  Table {ti}: {len(table.rows)} rows, "
                    f"{len(table.columns)} cols — "
                    f"headers: {headers}"
                )
            table_summary = "\n".join(table_lines)

        prompt = (
            f"The user wants to modify the existing template.\n"
            f"Current template content:\n{current_text[:3000]}\n\n"
            f"TABLE STRUCTURE (use these indices for modify_table "
            f"and add_table_loop):\n{table_summary}\n\n"
            f"User instruction: {instruction}\n\n"
            "Use your tools to apply the requested changes:\n"
            "- remove_section(heading_text) to remove an entire "
            "section (heading + content + renumber)\n"
            "- modify_table(table_index, remove_entire=True) to "
            "delete an entire table\n"
            "- modify_table(table_index, remove_columns=[...]) to "
            "remove specific columns\n"
            "- add_content(after_heading, paragraphs) to add "
            "formatted content (lists, text) after a heading. "
            "Use style 'List Number' for numbered lists, "
            "'List Bullet' for bullets, 'Normal' for plain text.\n"
            "- apply_substitutions to change variable placeholders\n"
            "- add_table_loop to restructure tables\n\n"
            "IMPORTANT:\n"
            "- To remove a section (heading + table/content): use "
            "remove_section which removes the heading AND all "
            "content until the next heading, then renumbers.\n"
            "- Use add_content to add structured content — do NOT "
            "cram multi-line text into apply_substitutions.\n"
            "- To convert a table to a list: first "
            "modify_table(remove_entire=True), then "
            "add_content(after_heading, paragraphs).\n\n"
            "AFTER making changes, you MUST call validate to verify "
            "the result. Check the validate output carefully — if "
            "there are unknown variables, missing steps loop, or "
            "other issues, fix them before finishing.\n\n"
            "QA CHECK: After validate passes, review the template "
            "text in the validate output. Verify that:\n"
            "1. The user's requested change was fully applied\n"
            "2. No orphaned headings, empty sections, or stale "
            "content remains\n"
            "3. Section numbering is sequential\n"
            "If something is incomplete, use more tools to fix it."
        )

        await asyncio.wait_for(
            agent.run(prompt, deps=deps),
            timeout=300,
        )

        # Build result
        detected_vars: set[str] = set()
        if state.exists("template.docx"):
            all_text = _extract_docx_text(state.read("template.docx"))
            detected_vars = _extract_jinja_variables(all_text)

        critical_vars = {"steps"}
        missing_critical = sorted(
            v for v in critical_vars if v not in detected_vars
        )

        warnings: list[dict[str, str]] = []
        for var in missing_critical:
            warnings.append({
                "type": "critical_missing",
                "variable": var,
                "description": (
                    f"CRITICAL: Template is missing '{var}' loop. "
                    "Every SOP/batch record template must have a "
                    "procedure steps loop."
                ),
            })

        missing = sorted(
            v for v in KNOWN_VARIABLES
            if v not in detected_vars and v not in critical_vars
        )
        for var in missing:
            warnings.append({
                "type": "missing_variable",
                "variable": var,
                "description": (
                    f"Template does not use '{var}' — this data "
                    "won't appear in the rendered output"
                ),
            })

        result_data = {
            "conversion_id": str(state.conversion_id),
            "preview_url": state.preview_url,
            "template_download_url": state.template_url,
            "warnings": warnings,
            "variables_detected": sorted(detected_vars),
        }

        _save_trace(
            deps, "refine", instruction[:100], str(model),
            "success",
        )

        event_stream.push("complete", {
            "template_url": state.template_url,
            "preview_url": (
                state.preview_url
                if state.exists("preview.pdf")
                else None
            ),
            "variables": sorted(detected_vars),
            "warnings": warnings,
        })

        return result_data

    except Exception as e:
        logger.exception("Template refinement failed")
        if deps:
            _save_trace(
                deps, "refine", instruction[:100], str(model),
                "error", error_msg=str(e),
            )
        event_stream.push("error", {"message": str(e)})
        raise
    finally:
        event_stream.close()


async def reupload_template(
    db: AsyncSession,
    org_id: UUID,
    state: ConversionState,
    file_bytes: bytes,
) -> dict[str, Any]:
    """Re-upload a manually edited template, re-render and re-validate."""
    state.write("template.docx", file_bytes)

    render_result = _try_render(file_bytes, "SOP")

    if render_result.pdf_bytes:
        state.write("preview.pdf", render_result.pdf_bytes)

    all_text = _extract_docx_text(file_bytes)
    detected_vars = _extract_jinja_variables(all_text)

    warnings: list[dict[str, str]] = []
    if render_result.render_error:
        warnings.append({
            "type": "render_error",
            "variable": "",
            "description": render_result.render_error,
        })
    if render_result.jinja_remnants:
        warnings.append({
            "type": "jinja_remnants",
            "variable": "",
            "description": (
                f"Raw Jinja2 syntax survived rendering: "
                f"{render_result.jinja_remnants}"
            ),
        })

    missing = sorted(v for v in KNOWN_VARIABLES if v not in detected_vars)
    for var in missing:
        warnings.append({
            "type": "missing_variable",
            "variable": var,
            "description": (
                f"Template does not use '{var}' — this data won't appear "
                "in the rendered output"
            ),
        })

    state.write_json("detected_vars.json", sorted(detected_vars))

    return {
        "conversion_id": str(state.conversion_id),
        "preview_url": state.preview_url,
        "template_download_url": state.template_url,
        "warnings": warnings,
        "variables_detected": sorted(detected_vars),
    }


async def save_to_library(
    db: AsyncSession,
    org_id: UUID,
    user_id: UUID,
    state: ConversionState,
    name: str,
    template_type: str,
    description: str | None,
    project_id: UUID | None,
    set_as_default: bool,
) -> dict[str, Any]:
    """Save the converted template to the DocumentTemplate library."""
    from app.models.templates import DocumentTemplate
    from app.services.template_engine import parse_template

    template_bytes = state.read("template.docx")

    storage = FileStorageService()
    filename = f"converted_{state.conversion_id}.docx"
    parts = [str(org_id), "document_templates", filename]
    relative_path = str(Path(*parts))
    full_path = storage.storage_root / relative_path
    full_path.parent.mkdir(parents=True, exist_ok=True)
    full_path.write_bytes(template_bytes)

    recognized, unrecognized = parse_template(full_path)

    template = DocumentTemplate(
        org_id=org_id,
        project_id=project_id,
        uploaded_by_id=user_id,
        name=name,
        description=description,
        template_type=template_type,
        file_path=relative_path,
        original_filename=filename,
        mime_type=(
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document"
        ),
        file_size_bytes=len(template_bytes),
        variables={
            "recognized": recognized,
            "unrecognized": unrecognized,
        },
    )
    db.add(template)
    await db.flush()

    return {
        "id": str(template.id),
        "name": template.name,
        "template_type": template.template_type,
        "status": template.status,
    }
