"""Build production-grade SOP + Batch Record .docx templates.

Run from repo root:
    python scripts/build_default_templates.py

Writes:
    backend/app/services/documents/templates/sop_default.docx
    backend/app/services/documents/templates/batch_record_default.docx

Templates use docxtpl Jinja2 syntax: {{ var }}, {% if %}, {% for %},
{%tr for %}, {%tc if %}. RichText tokens (e.g. {{r step.sop_body }}) are
rendered as paragraph runs that already contain newlines and page breaks
when build_context populates them.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.table import _Cell
from docx.text.paragraph import Paragraph

REPO_ROOT = Path(__file__).resolve().parent.parent
TEMPLATES_DIR = REPO_ROOT / "backend/app/services/documents/templates"

# SOP design tokens — refined scientific editorial. Cambria for prose,
# Calibri for labels and the procedure number column. Grayscale ink only.
INK = "1F2937"           # body text
INK_DEEP = "0F172A"      # headings, step names
INK_MUTED = "64748B"     # meta, captions
RULE = "94A3B8"          # paragraph borders
SHADE = "F1F5F9"         # table header fill
SERIF = "Cambria"
SANS = "Calibri"


def _p(
    doc: DocxDocument,
    text: str,
    bold: bool = False,
    size: int = 11,
    align: WD_ALIGN_PARAGRAPH | None = None,
) -> Paragraph:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


# ── SOP typography helpers ────────────────────────────────────────────


def _set_pBdr(
    p: Paragraph,
    position: str,
    color: str,
    *,
    size: int = 6,
    space: int = 4,
) -> None:
    """Apply a paragraph border edge. `size` is eighths of a point
    (6 = 0.75pt); `space` is the gap between border and text in pt."""
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    edge = OxmlElement(f"w:{position}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), str(space))
    edge.set(qn("w:color"), color)
    pBdr.append(edge)


def _set_letter_spacing(run, twentieths: int) -> None:
    """Track-out runs by `twentieths` of a point. 40 = 2pt letter spacing."""
    rPr = run._element.get_or_add_rPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), str(twentieths))
    rPr.append(spacing)


def _shade_cell(cell: _Cell, color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def _set_sop_default_style(doc: DocxDocument) -> None:
    """Cambria 11pt body, dark slate ink."""
    style = doc.styles["Normal"]
    style.font.name = SERIF
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor.from_string(INK)


def _section_heading(doc: DocxDocument, text: str) -> Paragraph:
    """Calibri 12pt bold, uppercase, letter-spaced, with thin bottom
    rule. Generous space before, kept with next so the heading never
    strands at the bottom of a page."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(22)
    pf.space_after = Pt(8)
    pf.keep_with_next = True
    run = p.add_run(text.upper())
    run.font.name = SANS
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(INK_DEEP)
    _set_letter_spacing(run, 40)  # 2pt tracking
    _set_pBdr(p, "bottom", RULE, size=4, space=6)
    return p


def _doc_subtitle(doc: DocxDocument, text: str) -> Paragraph:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = SANS
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(INK_MUTED)
    _set_letter_spacing(run, 60)
    return p


def _doc_title(doc: DocxDocument, text: str) -> Paragraph:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = SERIF
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(INK_DEEP)
    return p


def _doc_meta(doc: DocxDocument, text: str) -> Paragraph:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = SANS
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(INK_MUTED)
    return p


def _body_p(doc: DocxDocument, text: str) -> Paragraph:
    """Body paragraph — 11pt Cambria, 1.35 line height, space-after 6pt."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.35
    run = p.add_run(text)
    run.font.name = SERIF
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(INK)
    return p


def _control_p(doc: DocxDocument, text: str) -> Paragraph:
    """Tiny paragraph carrying a Jinja control tag. Minimised so the
    leftover whitespace after docxtpl strips the tag is negligible."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(1)
    return p


def _table_header_cell(cell: _Cell, text: str) -> None:
    """Style a table header cell: Calibri bold 9pt all-caps, shaded."""
    _shade_cell(cell, SHADE)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.text = ""
    run = p.add_run(text.upper())
    run.font.name = SANS
    run.font.size = Pt(9)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(INK_DEEP)
    _set_letter_spacing(run, 30)


def _add_role_header_paragraph(doc: DocxDocument) -> Paragraph:
    """Template paragraph that the RichText role header lands in.
    Carries a bottom border and generous spacing; the RichText itself
    sets fonts/sizes/colors for the role name."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(10)
    pf.keep_with_next = True
    _set_pBdr(p, "bottom", RULE, size=4, space=6)
    return p


def _add_procedure_step_paragraph(doc: DocxDocument) -> Paragraph:
    """Template paragraph for a single procedure step. Hanging indent
    at 0.4" so the step number sits in the left margin and wrapped
    lines align under the step text. Tab stop at the indent so the
    \\t between number and step name lands flush."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.4)
    pf.first_line_indent = Inches(-0.4)
    pf.space_after = Pt(10)
    pf.line_spacing = 1.3
    pf.keep_together = True
    pf.tab_stops.add_tab_stop(Inches(0.4))
    return p


def _table_loop(
    table,
    header_cells: list[str],
    iterator: str,
    data_cells: list[str],
) -> None:
    """Wire up a docxtpl-style table loop. Header row gets styled cells;
    a {%tr for %} row, a data row, and a {%tr endfor %} row follow."""
    for cell, label in zip(table.rows[0].cells, header_cells):
        _table_header_cell(cell, label)
    for_row = table.add_row().cells
    for_row[0].text = iterator
    data_row = table.add_row().cells
    for cell, value in zip(data_row, data_cells):
        cell.text = value
        # Body cell typography
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = SERIF
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor.from_string(INK)
    end_row = table.add_row().cells
    end_row[0].text = "{%tr endfor %}"


def build_sop(output_path: Path) -> None:
    """Production-grade SOP template — refined scientific editorial.

    Cambria body + Calibri labels, grayscale ink, real Word paragraph
    borders for section dividers (no Unicode rules), hanging-indent
    numbered procedure with a tab-aligned number column.
    """
    doc = Document()
    _set_sop_default_style(doc)

    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)

    # Page header — small uppercase tag with org / doc # / version.
    hdr = section.header.paragraphs[0]
    hdr.text = ""
    hdr_run = hdr.add_run(
        "{{ organization_name }}    "
        "{% if doc_number %}{{ doc_number }}    {% endif %}"
        "Version {{ version_number }}"
    )
    hdr_run.font.name = SANS
    hdr_run.font.size = Pt(8.5)
    hdr_run.font.color.rgb = RGBColor.from_string(INK_MUTED)
    _set_letter_spacing(hdr_run, 30)

    # Title block: tracked-out kicker, large serif title, meta sub-line,
    # project/org footer, then a hairline rule that separates the title
    # block from the document body.
    _doc_subtitle(
        doc,
        "{% if doc_number %}{{ doc_number }}  ·  {% endif %}"
        "STANDARD OPERATING PROCEDURE",
    )
    _doc_title(doc, "{{ protocol_name }}")
    _doc_meta(
        doc,
        "Version {{ version_number }}  ·  "
        "Effective {{ effective_date }}  ·  "
        "Supersedes {{ supersedes_date }}",
    )
    _doc_meta(doc, "{{ project_name }}  ·  {{ organization_name }}")
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(6)
    rule.paragraph_format.space_after = Pt(4)
    rule.add_run("").font.size = Pt(1)
    _set_pBdr(rule, "bottom", RULE, size=4, space=2)

    # Gated text sections — Purpose / Scope / Definitions / References.
    for heading, gate, body in [
        ("Purpose", "purpose", "{{ purpose }}"),
        ("Scope", "scope", "{{ scope }}"),
        ("Definitions", "definitions", "{{ definitions }}"),
        ("References", "references", "{{ references }}"),
    ]:
        _control_p(doc, "{%p if " + gate + " %}")
        _section_heading(doc, heading)
        _body_p(doc, body)
        _control_p(doc, "{%p endif %}")

    # Revision History table.
    _control_p(doc, "{%p if revision_history %}")
    _section_heading(doc, "Revision History")
    _table_loop(
        doc.add_table(rows=1, cols=4),
        header_cells=["Version", "Date", "Author", "Changes"],
        iterator="{%tr for rev in revision_history %}",
        data_cells=[
            "{{ rev.version_number }}",
            "{{ rev.created_at }}",
            "{{ rev.created_by }}",
            "{{ rev.change_summary }}",
        ],
    )
    _control_p(doc, "{%p endif %}")

    # Responsibilities table.
    _control_p(doc, "{%p if responsibilities %}")
    _section_heading(doc, "Responsibilities")
    _table_loop(
        doc.add_table(rows=1, cols=2),
        header_cells=["Role", "Responsibility summary"],
        iterator="{%tr for resp in responsibilities %}",
        data_cells=["{{ resp.role_name }}", "{{ resp.step_summary }}"],
    )
    _control_p(doc, "{%p endif %}")

    # Equipment summary table.
    _control_p(doc, "{%p if equipment_summary %}")
    _section_heading(doc, "Equipment")
    _table_loop(
        doc.add_table(rows=1, cols=3),
        header_cells=["ID", "Name", "Description"],
        iterator="{%tr for eq in equipment_summary %}",
        data_cells=[
            "{{ eq.local_id }}",
            "{{ eq.name }}",
            "{{ eq.description }}",
        ],
    )
    _control_p(doc, "{%p endif %}")

    # Procedure — always present, branches on is_role_based. The step
    # number, name, description, and meta lines are pre-composed in
    # template_engine._build_sop_body so the RichText carries its own
    # typography; the template paragraphs only contribute layout
    # (hanging indent + tab stop) and the role-header divider rule.
    _section_heading(doc, "Procedure")
    _control_p(doc, "{%p if is_role_based %}")
    _control_p(doc, "{%p for role in roles %}")
    role_hdr = _add_role_header_paragraph(doc)
    role_hdr.add_run("{{r role.sop_header }}")
    _control_p(doc, "{%p for step in role.sop_steps %}")
    step_p = _add_procedure_step_paragraph(doc)
    step_p.add_run("{{r step.sop_body }}")
    _control_p(doc, "{%p endfor %}")
    _control_p(doc, "{%p endfor %}")
    _control_p(doc, "{%p else %}")
    _control_p(doc, "{%p for step in steps %}")
    flat_step_p = _add_procedure_step_paragraph(doc)
    flat_step_p.add_run("{{r step.sop_body }}")
    _control_p(doc, "{%p endfor %}")
    _control_p(doc, "{%p endif %}")

    # Approval block.
    _section_heading(doc, "Approval")
    _control_p(doc, "{%p if approval %}")
    _body_p(
        doc,
        "Approved by: {{ approval.actor_name }} "
        "({{ approval.actor_role }})",
    )
    _body_p(doc, "Date: {{ approval.approved_at }}")
    _control_p(doc, "{%p if approval.signature_statement %}")
    _body_p(doc, "Statement: {{ approval.signature_statement }}")
    _control_p(doc, "{%p endif %}")
    _control_p(doc, "{%p else %}")
    _body_p(doc, "{{ unapproved_warning }}")
    _control_p(doc, "{%p endif %}")

    doc.save(output_path)


# ── Batch Record-specific helpers ─────────────────────────────────────


def _clear_table_borders(table) -> None:
    """Strip the default Word grid so the table reads as borderless
    typography (used for the run-details key/value block)."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = tblBorders.find(qn(f"w:{edge}"))
        if b is None:
            b = OxmlElement(f"w:{edge}")
            tblBorders.append(b)
        b.set(qn("w:val"), "nil")


def _indent_table(table, indent: Inches) -> None:
    """Indent the entire table by `indent` so it visually aligns with
    the step-card body paragraphs (which sit at the same indent)."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    # `dxa` units = 20ths of a point = twips
    tblInd.set(qn("w:w"), str(int(indent.twips)))
    tblInd.set(qn("w:type"), "dxa")


def _br_kv_row(row, label: str, value_expr: str) -> None:
    """Render one row of the run-details key/value block.  Left cell:
    uppercase letter-spaced Calibri muted label.  Right cell: Cambria
    serif value, where `value_expr` is a Jinja expression."""
    lp = row.cells[0].paragraphs[0]
    lp.text = ""
    lp.paragraph_format.space_after = Pt(3)
    lrun = lp.add_run(label.upper())
    lrun.font.name = SANS
    lrun.font.size = Pt(8.5)
    lrun.font.color.rgb = RGBColor.from_string(INK_MUTED)
    _set_letter_spacing(lrun, 40)
    vp = row.cells[1].paragraphs[0]
    vp.text = ""
    vp.paragraph_format.space_after = Pt(3)
    vrun = vp.add_run(value_expr)
    vrun.font.name = SERIF
    vrun.font.size = Pt(10.5)
    vrun.font.color.rgb = RGBColor.from_string(INK)


def _br_run_details(doc: DocxDocument, rows: list[tuple[str, str]]) -> None:
    """Borderless 2-column key/value table for the run-details panel."""
    t = doc.add_table(rows=len(rows), cols=2)
    t.autofit = False
    _clear_table_borders(t)
    for (label, value), row in zip(rows, t.rows):
        row.cells[0].width = Inches(1.8)
        row.cells[1].width = Inches(4.5)
        _br_kv_row(row, label, value)


def _br_caption(doc: DocxDocument, text: str) -> Paragraph:
    """Small uppercase letter-spaced caption inside a step card —
    `PARAMETERS`, `EXECUTION`, `NOTES`. Indented to match the card body
    so it reads as a sub-label of the step, not a section heading."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.4)
    pf.space_before = Pt(8)
    pf.space_after = Pt(2)
    pf.keep_with_next = True
    run = p.add_run(text.upper())
    run.font.name = SANS
    run.font.size = Pt(8.5)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(INK_MUTED)
    _set_letter_spacing(run, 40)
    return p


def _br_step_header_paragraph(doc: DocxDocument) -> Paragraph:
    """First paragraph of a step card.  A top hairline rule visually
    separates each card from the previous one.  Tab stops at 0.4in
    (number column) and ~5.8in right-aligned (duration column) match
    the layout baked into `_build_br_card_header` on the engine side."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.4)
    pf.first_line_indent = Inches(-0.4)
    pf.space_before = Pt(14)
    pf.space_after = Pt(2)
    pf.keep_with_next = True
    pf.tab_stops.add_tab_stop(Inches(0.4))
    pf.tab_stops.add_tab_stop(Inches(5.8), WD_TAB_ALIGNMENT.RIGHT)
    _set_pBdr(p, "top", RULE, size=4, space=8)
    return p


def _br_indented_body_p(doc: DocxDocument, text: str) -> Paragraph:
    """Step-card body paragraph (description, notes).  Indented 0.4in
    so it aligns under the step name, smaller leading than the SOP body
    because step cards stack denser than flowing prose."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.4)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.3
    run = p.add_run(text)
    run.font.name = SERIF
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(INK)
    return p


def _br_value_paragraph(doc: DocxDocument) -> Paragraph:
    """Empty paragraph the `{{r step.value_display }}` RichText lands
    in.  The RichText itself sets fonts/sizes."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.4)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.3
    return p


def _br_exec_data_cell(cell: _Cell, value_expr: str) -> None:
    """Style an execution-row data cell — Cambria 10pt, ink body."""
    p = cell.paragraphs[0]
    p.text = ""
    run = p.add_run(value_expr)
    run.font.name = SERIF
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(INK)


def _emit_br_step_card(doc: DocxDocument) -> None:
    """Render one Batch Record step card.

    Anatomy (Jinja loop body — repeated per step):
        [top hairline]
        [step header paragraph]   ← {{r step.card_header }}
        [description paragraph]   ← {{ step.description }} (conditional)
        [PARAMETERS caption]
        [value_display paragraph] ← {{r step.value_display }}
        [EXECUTION caption]
        [5-col execution table]   ← initials, reviewer, scheduled, start, end
        [NOTES caption]           ← (conditional)
        [notes paragraph]         ← (conditional)
    """
    # Step header — number + name + duration
    sh = _br_step_header_paragraph(doc)
    sh.add_run("{{r step.card_header }}")

    # Description — skip the whole paragraph if empty
    _control_p(doc, "{%p if step.description %}")
    _br_indented_body_p(doc, "{{ step.description }}")
    _control_p(doc, "{%p endif %}")

    # Parameters caption + value display
    _br_caption(doc, "Parameters")
    vp = _br_value_paragraph(doc)
    vp.add_run("{{r step.value_display }}")

    # Execution caption + 5-column mini table
    _br_caption(doc, "Execution")
    exec_tbl = doc.add_table(rows=2, cols=5)
    _indent_table(exec_tbl, Inches(0.4))
    headers = ["Operator", "Reviewer", "Scheduled", "Started", "Completed"]
    for cell, label in zip(exec_tbl.rows[0].cells, headers):
        _table_header_cell(cell, label)
    data = exec_tbl.rows[1].cells
    _br_exec_data_cell(data[0], "{{ step.initials }}")
    _br_exec_data_cell(
        data[1], "{{ step.reviewer_initials if reviewer_enabled else '' }}",
    )
    _br_exec_data_cell(
        data[2], "{{ step.scheduled_at if time_enabled else '' }}",
    )
    _br_exec_data_cell(
        data[3], "{{ step.actual_started_at if time_enabled else '' }}",
    )
    _br_exec_data_cell(
        data[4], "{{ step.actual_completed_at if time_enabled else '' }}",
    )

    # Notes — caption + paragraph, both inside a single conditional
    _control_p(doc, "{%p if step.notes_text %}")
    _br_caption(doc, "Notes")
    _br_indented_body_p(doc, "{{ step.notes_text }}")
    _control_p(doc, "{%p endif %}")


def build_batch_record(output_path: Path) -> None:
    """Production-grade Batch Record template — refined scientific
    editorial, adapted for journalistic (data-capture) rhythm.

    Where the SOP is editorial prose with a numbered procedure list, the
    BR is a sequence of step *cards*: each card holds the executed values
    for one step (parameters / operator / reviewer / timestamps / notes).
    Equipment, deviations, and approval surround the procedure with the
    same uppercase letter-spaced section headings as the SOP.
    """
    doc = Document()
    _set_sop_default_style(doc)  # Cambria 11pt Normal — shared with SOP

    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)

    # Page header — letter-spaced Calibri tag with org / run / lot.
    hdr = section.header.paragraphs[0]
    hdr.text = ""
    hdr_run = hdr.add_run(
        "{{ organization_name }}    "
        "BATCH RECORD    "
        "Run: {{ run_name }}    "
        "{% if lot_number %}Lot {{ lot_number }}{% endif %}"
    )
    hdr_run.font.name = SANS
    hdr_run.font.size = Pt(8.5)
    hdr_run.font.color.rgb = RGBColor.from_string(INK_MUTED)
    _set_letter_spacing(hdr_run, 30)

    # Title block — same rhythm as SOP: tracked-out kicker, 24pt serif
    # title, muted meta lines, hairline rule. Lot/Batch lines hang off
    # the meta block as their own subdued paragraphs so they're
    # searchable and visually grouped under the title.
    _doc_subtitle(doc, "BATCH RECORD")
    _doc_title(doc, "{{ protocol_name }}")
    _doc_meta(
        doc,
        "Run {{ run_name }}  ·  Status {{ run_status }}  ·  "
        "Started {{ started_at }}",
    )
    _doc_meta(doc, "{{ project_name }}  ·  {{ organization_name }}")
    _control_p(doc, "{%p if lot_number %}")
    _doc_meta(doc, "Lot Number: {{ lot_number }}")
    _control_p(doc, "{%p endif %}")
    _control_p(doc, "{%p if batch_number %}")
    _doc_meta(doc, "Batch Number: {{ batch_number }}")
    _control_p(doc, "{%p endif %}")
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(6)
    rule.paragraph_format.space_after = Pt(4)
    rule.add_run("").font.size = Pt(1)
    _set_pBdr(rule, "bottom", RULE, size=4, space=2)

    # Run Details — borderless 2-col key/value block. Reads like a
    # masthead under the title.
    _section_heading(doc, "Run Details")
    _br_run_details(doc, [
        ("Protocol",
         "{{ protocol_name }} (v{{ version_number }})"),
        ("SOP reference",
         "{{ doc_number }}  ·  effective {{ effective_date }}"),
        ("Project", "{{ project_name }}"),
        ("Organization", "{{ organization_name }}"),
        ("Run name", "{{ run_name }}"),
        ("Status", "{{ run_status }}"),
        ("Started", "{{ started_at }}"),
        ("Completed", "{{ completed_at }}"),
    ])

    # Equipment Used — SOP-equipment table shape; uppercase heading
    # legible enough that auditors can scan it.
    _control_p(doc, "{%p if equipment_summary %}")
    _section_heading(doc, "Equipment Used")
    _table_loop(
        doc.add_table(rows=1, cols=3),
        header_cells=["ID", "Name", "Description"],
        iterator="{%tr for eq in equipment_summary %}",
        data_cells=[
            "{{ eq.local_id }}",
            "{{ eq.name }}",
            "{{ eq.description }}",
        ],
    )
    _control_p(doc, "{%p endif %}")

    # Execution Record — step cards (one per step). The role-based
    # branch wraps the inner step loop in `{%p for role %}` and prints
    # a role-header paragraph between roles; the flat branch just
    # iterates `steps` directly.
    _section_heading(doc, "Execution Record")
    _control_p(doc, "{%p if is_role_based %}")
    _control_p(doc, "{%p for role in roles %}")
    role_hdr = _add_role_header_paragraph(doc)
    role_hdr.add_run("{{r role.br_header }}")
    _control_p(doc, "{%p for step in role.steps %}")
    _emit_br_step_card(doc)
    _control_p(doc, "{%p endfor %}")
    _control_p(doc, "{%p endfor %}")
    _control_p(doc, "{%p else %}")
    _control_p(doc, "{%p for step in steps %}")
    _emit_br_step_card(doc)
    _control_p(doc, "{%p endfor %}")
    _control_p(doc, "{%p endif %}")

    # Deviations Log — appears only when something was flagged.
    _control_p(doc, "{%p if deviations %}")
    _section_heading(doc, "Deviations Log")
    _table_loop(
        doc.add_table(rows=1, cols=3),
        header_cells=["When", "Author", "Note"],
        iterator="{%tr for d in deviations %}",
        data_cells=[
            "{{ d.created_at }}",
            "{{ d.author_name }}",
            "{{ d.content }}",
        ],
    )
    _control_p(doc, "{%p endif %}")

    # Run Notes — author·timestamp caption, then the note body.
    _control_p(doc, "{%p if notes %}")
    _section_heading(doc, "Run Notes")
    _control_p(doc, "{%p for note in notes %}")
    _doc_meta(doc, "{{ note.author_name }}  ·  {{ note.created_at }}")
    _body_p(doc, "{{ note.content }}")
    _control_p(doc, "{%p endfor %}")
    _control_p(doc, "{%p endif %}")

    # Figures
    _control_p(doc, "{%p if figures %}")
    _section_heading(doc, "Figures")
    _control_p(doc, "{%p for fig in figures %}")
    _doc_meta(doc, "Figure {{ fig.number }}: {{ fig.filename }}")
    _body_p(doc, "{{ fig.image }}")
    _control_p(doc, "{%p endfor %}")
    _control_p(doc, "{%p endif %}")

    # Attachments (non-image)
    _control_p(doc, "{%p if non_image_attachments %}")
    _section_heading(doc, "Attachments")
    _control_p(doc, "{%p for att in non_image_attachments %}")
    _body_p(doc, "{{ att.filename }}  ·  {{ att.uploaded_at }}")
    _control_p(doc, "{%p endfor %}")
    _control_p(doc, "{%p endif %}")

    # Approval — labeled signature/date rows; unapproved-warning when
    # approval is missing. Approval history (if multi-event) trails as
    # a small italic appendix.
    _section_heading(doc, "Approval")
    _control_p(doc, "{%p if approval %}")
    _body_p(
        doc,
        "Approved by: {{ approval.actor_name }} "
        "({{ approval.actor_role }})",
    )
    _body_p(doc, "Date: {{ approval.approved_at }}")
    _control_p(doc, "{%p if approval.signature_statement %}")
    _body_p(doc, "Statement: {{ approval.signature_statement }}")
    _control_p(doc, "{%p endif %}")
    _control_p(doc, "{%p else %}")
    _body_p(doc, "{{ unapproved_warning }}")
    _control_p(doc, "{%p endif %}")
    _control_p(doc, "{%p if approval_history and approval_history|length > 1 %}")
    _doc_meta(doc, "Approval history:")
    _control_p(doc, "{%p for event in approval_history %}")
    _doc_meta(
        doc,
        "{{ event.action }} by {{ event.actor_name }} "
        "on {{ event.created_at }}",
    )
    _control_p(doc, "{%p endfor %}")
    _control_p(doc, "{%p endif %}")

    doc.save(output_path)


if __name__ == "__main__":
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    sop = TEMPLATES_DIR / "sop_default.docx"
    br = TEMPLATES_DIR / "batch_record_default.docx"
    build_sop(sop)
    build_batch_record(br)
    print(f"wrote {sop}")
    print(f"wrote {br}")
