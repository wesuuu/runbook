"""Inject the F-0066 approval block into the default SOP and Batch Record
docx templates.

Idempotent — running again is a no-op (looks for an "approval-marker"
sentinel paragraph and skips when already present).

Usage:
    python scripts/inject_approval_section.py

The Jinja blocks added here are evaluated by docxtpl using the context
keys ``unapproved_warning``, ``approval``, and ``approval_history``
populated by ``_build_approval_context`` in the PDF endpoints.

Note: signature image embedding requires the docxtpl ``InlineImage``
helper, which is wired up via render_to_docx (see template_engine.py).
``approval.signature_image`` is rendered to an InlineImage there when
``approval.signature_image_path`` is present.
"""

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

MARKER = "BR-APPROVAL-BLOCK"  # sentinel comment we leave in a hidden run

WARNING_BLOCK = [
    "{%p if unapproved_warning %}",
    "⚠ UNAPPROVED — DRAFT ONLY",
    "{%p endif %}",
]

# Renders an "Approval & Signatures" section after the body. The
# approval.signature_image is swapped to an InlineImage by render_to_docx
# when a signature_image_path is present on the approval dict.
APPROVAL_SECTION = [
    f"{{# {MARKER} #}}",
    "{%p if approval %}",
    "Approval & Signatures",
    "Approver: {{ approval.approver_name }} <{{ approval.approver_email }}>",
    "Approved at: {{ approval.approved_at }}    Protocol version: {{ approval.protocol_version }}",
    "{%p if approval.signature_statement %}",
    "Statement: {{ approval.signature_statement }}",
    "{%p endif %}",
    "{%p if approval.signature_image %}",
    "Signature: {{ approval.signature_image }}",
    "{%p endif %}",
    "{%p endif %}",
    "{%p if approval_history %}",
    "Approval History",
    "{%p for ev in approval_history %}",
    "{{ ev.created_at }} — {{ ev.action }} by {{ ev.actor_name }}",
    "{%p endfor %}",
    "{%p endif %}",
]


def _already_injected(doc) -> bool:
    for p in doc.paragraphs:
        if MARKER in p.text:
            return True
    return False


def _insert_paragraphs_at_start(doc, lines: list[str]) -> None:
    """Insert plain paragraphs at the top of the document body."""
    body = doc.element.body
    sectpr = body.find(qn("w:sectPr"))
    # Anchor for insertion: the very first child of body
    first = body[0]
    # We add new paragraphs to the doc, then move them before `first`.
    new_paras = []
    for line in lines:
        p = doc.add_paragraph(line)
        new_paras.append(p._p)
    # add_paragraph appends to body (before sectPr); move each before
    # `first` in original order — addprevious places each one
    # immediately before the anchor, so iterating forward yields
    # correct top-to-bottom order in the final document.
    for p_el in new_paras:
        body.remove(p_el)
        first.addprevious(p_el)


def _append_paragraphs(doc, lines: list[str]) -> None:
    for line in lines:
        doc.add_paragraph(line)


def inject(template_path: Path) -> bool:
    doc = Document(str(template_path))
    if _already_injected(doc):
        return False
    _insert_paragraphs_at_start(doc, WARNING_BLOCK)
    _append_paragraphs(doc, APPROVAL_SECTION)
    doc.save(str(template_path))
    return True


def main():
    base = Path(__file__).resolve().parent.parent / (
        "backend/app/services/documents/templates"
    )
    for name in ("sop_default.docx", "batch_record_default.docx"):
        path = base / name
        changed = inject(path)
        print(f"{name}: {'updated' if changed else 'already injected'}")


if __name__ == "__main__":
    main()
